import os
import re
import sqlite3
import random
import time
import requests
import markdown as md
import bleach
from functools import wraps
from io import BytesIO
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash, abort
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from htmldocx import HtmlToDocx
from xhtml2pdf import pisa
from bs4 import BeautifulSoup

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "chave_mestra_professor_ia_2026")
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# 1. MAPEAMENTO DAS CHAVES DE IA
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
# Chave opcional do Mistral — usada apenas como reforço extra (gratuito) caso
# os modelos do Gemini falhem em sequência. Se não configurada, é ignorada.
MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "").strip()

# Dicionário de módulos unificado
MODULOS = {
    'plano': {'nome': 'Plano de Aula', 'icone': 'fa-book'},
    'bimestral': {'nome': 'Planejamento Bimestral', 'icone': 'fa-calendar-check'},
    'atividades': {'nome': 'Banco de Atividades', 'icone': 'fa-list-check'},
    'avaliacoes': {'nome': 'Gerador de Provas', 'icone': 'fa-file-signature'},
    'simulados': {'nome': 'Simulados', 'icone': 'fa-clipboard-question'},
    'sequencia': {'nome': 'Sequência Didática', 'icone': 'fa-layer-group'},
    'duvidas': {'nome': 'Tira-Dúvidas com IA', 'icone': 'fa-circle-question'},
    'relatorios': {'nome': 'Relatórios Pedagógicos', 'icone': 'fa-chart-line'},
    'inclusao': {'nome': 'Plano de Inclusão / AEE', 'icone': 'fa-hands-asl-interpreting'},
    'projetos': {'nome': 'Projetos Interdisciplinares', 'icone': 'fa-diagram-project'},
    'alfabetizacao': {'nome': 'Alfabetização e Reforço de Leitura', 'icone': 'fa-spell-check'}
}

TIPOS_SIMULADO = [
    "Simulado Bimestral / Revisão de Conteúdo",
    "Simulado tipo ENEM",
    "Simulado tipo Vestibular",
    "Simulado SAEB / Prova Brasil",
    "Simulado de Concurso Público",
    "Simulado Geral / Diagnóstico"
]

NIVEIS_LEITURA = [
    "Pré-silábico (não reconhece letras/sons)",
    "Silábico sem valor sonoro (junta símbolos sem som correspondente)",
    "Silábico com valor sonoro (associa algumas letras ao som)",
    "Silábico-alfabético (mistura sílabas e letras corretamente)",
    "Alfabético com dificuldades (lê, mas troca/omite letras, lê devagar)",
    "Alfabético fluente com dificuldade de interpretação (lê bem, mas não compreende)"
]

FOCOS_ALFABETIZACAO = [
    "Consciência Fonológica (rimas, sons, sílabas)",
    "Reconhecimento de Letras (alfabeto, maiúsculas/minúsculas)",
    "Formação de Sílabas",
    "Leitura de Palavras",
    "Leitura de Frases e Textos Curtos",
    "Fluência Leitora (velocidade e ritmo)",
    "Compreensão Leitora (interpretação)",
    "Escrita Espontânea",
    "Ortografia Básica",
    "Gosto pela Leitura / Motivação"
]

TIPOS_ATIVIDADE = [
    "Interpretação de texto", "Questões objetivas", "Questões subjetivas", "Produção textual",
    "Atividade de leitura (alfabetização)", "Complete as lacunas", "Ligue as colunas",
    "Verdadeiro ou falso", "Caça-palavras", "Cruzadinha", "Sequência lógica",
    "Ordenação de frases", "Escrita de palavras", "Formação de frases", "Problemas matemáticos",
    "Atividade ilustrada", "Recorte e cole", "Pintura educativa", "Pesquisa", "Debate em sala",
    "Trabalho em grupo", "Atividade prática", "Revisão de conteúdo", "Exercícios para reforço"
]

TIPOS_PROJETO = [
    "Feira de Ciências", "Feira Cultural", "Feira Literária", "Mostra Pedagógica", "Visita Técnica",
    "Aula de Campo", "Projeto Interdisciplinar", "Projeto de Leitura", "Projeto de Escrita",
    "Projeto Ambiental", "Projeto de Sustentabilidade", "Projeto de História Local",
    "Projeto de Cultura Popular", "Projeto de Consciência Negra", "Projeto de Educação Financeira",
    "Projeto de Saúde", "Projeto Esportivo", "Gincana Educativa", "Culminância",
    "Exposição de Trabalhos", "Oficina", "Seminário", "Palestra", "Teatro", "Musical",
    "Produção de Podcast", "Produção de Vídeo", "Robótica Educacional", "Clube de Ciências",
    "Clube de Leitura", "Horta Escolar", "Outro"
]

# =====================================================================
# TEXTO OFICIAL FIXO — COMPETÊNCIAS GERAIS DA EDUCAÇÃO BÁSICA (BNCC, 2018)
# Fixado no backend (não gerado pela IA) para garantir fidelidade 100% ao
# texto oficial do MEC em todo Planejamento Bimestral, independente do tema.
# =====================================================================
COMPETENCIAS_GERAIS_BNCC = [
    "Valorizar e utilizar os conhecimentos historicamente construídos sobre o mundo físico, social, cultural e digital para entender e explicar a realidade, continuar aprendendo e colaborar para a construção de uma sociedade justa, democrática e inclusiva.",
    "Exercitar a curiosidade intelectual e recorrer à abordagem própria das ciências, incluindo a investigação, a reflexão, a análise crítica, a imaginação e a criatividade, para investigar causas, elaborar e testar hipóteses, formular e resolver problemas e criar soluções (inclusive tecnológicas) com base nos conhecimentos das diferentes áreas.",
    "Valorizar e fruir as diversas manifestações artísticas e culturais, das locais às mundiais, e também participar de práticas diversificadas da produção artístico-cultural.",
    "Utilizar diferentes linguagens – verbal (oral ou visual-motora, como Libras, e escrita), corporal, visual, sonora e digital –, bem como conhecimentos das linguagens artística, matemática e científica, para se expressar e partilhar informações, experiências, ideias e sentimentos em diferentes contextos e produzir sentidos que levem ao entendimento mútuo.",
    "Compreender, utilizar e criar tecnologias digitais de informação e comunicação de forma crítica, significativa, reflexiva e ética nas diversas práticas sociais (incluindo as escolares) para se comunicar, acessar e disseminar informações, produzir conhecimentos, resolver problemas e exercer protagonismo e autoria na vida pessoal e coletiva.",
    "Valorizar a diversidade de saberes e vivências culturais e apropriar-se de conhecimentos e experiências que lhe possibilitem entender as relações próprias do mundo do trabalho e fazer escolhas alinhadas ao exercício da cidadania e ao seu projeto de vida, com liberdade, autonomia, consciência crítica e responsabilidade.",
    "Argumentar com base em fatos, dados e informações confiáveis, para formular, negociar e defender ideias, pontos de vista e decisões comuns que respeitem e promovam os direitos humanos, a consciência socioambiental e o consumo responsável em âmbito local, regional e global, com posicionamento ético em relação ao cuidado de si mesmo, dos outros e do planeta.",
    "Conhecer-se, apreciar-se e cuidar de sua saúde física e emocional, compreendendo-se na diversidade humana e reconhecendo suas emoções e as dos outros, com autocrítica e capacidade para lidar com elas.",
    "Exercitar a empatia, o diálogo, a resolução de conflitos e a cooperação, fazendo-se respeitar e promovendo o respeito ao outro e aos direitos humanos, com acolhimento e valorização da diversidade de indivíduos e de grupos sociais, seus saberes, identidades, culturas e potencialidades, sem preconceitos de qualquer natureza.",
    "Agir pessoal e coletivamente com autonomia, responsabilidade, flexibilidade, resiliência e determinação, tomando decisões com base em princípios éticos, democráticos, inclusivos, sustentáveis e solidários.",
]

def montar_prova_duas_colunas(html_questoes):
    """
    Recebe o HTML das questões (cada uma em <div class="questao-item">) e
    monta um bloco de texto em 2 colunas reais (CSS column-count), fonte
    Arial 12 — sem o gabarito.

    Retorna uma TUPLA (bloco_colunas_html, gabarito_html). O gabarito NUNCA
    é incluído no bloco impresso/exportado — ele é devolvido separadamente
    para ser exibido apenas na aba "Informações Pedagógicas" (tela do
    professor), nunca na prova/simulado do aluno.

    IMPORTANTE: usamos column-count (fluxo de texto tipo "jornal") em vez de
    uma tabela HTML de 2 colunas fixas. Uma tabela força uma divisão rígida
    de questões entre as colunas; quando uma questão é maior que as outras,
    o navegador precisa manter a "linha" da tabela inteira, criando vãos
    enormes na impressão. Com column-count, o texto flui naturalmente e
    quebra de página corretamente.
    """
    soup = BeautifulSoup(html_questoes, 'html.parser')
    itens = soup.find_all('div', class_='questao-item')
    gabarito_tag = soup.find('div', class_='gabarito-prova')
    gabarito_html = str(gabarito_tag) if gabarito_tag else ''

    fonte_estilo = 'font-family: Arial, Helvetica, sans-serif; font-size: 12pt; line-height: 1.4;'

    if not itens:
        # Rede de segurança: a IA não usou as divs esperadas — devolve sem split, mas já com a fonte correta.
        conteudo_sem_gabarito = html_questoes
        if gabarito_tag:
            conteudo_sem_gabarito = conteudo_sem_gabarito.replace(str(gabarito_tag), '')
        return f'<div style="{fonte_estilo}">{conteudo_sem_gabarito}</div>', gabarito_html

    questoes_html = "".join(str(i) for i in itens)

    bloco_colunas = f"""
    <div class="prova-colunas" style="column-count:2; -webkit-column-count:2; column-gap:30px; -webkit-column-gap:30px; column-rule:1px solid #999; {fonte_estilo}">
        {questoes_html}
    </div>
    """
    return bloco_colunas, gabarito_html


def montar_html_competencias_gerais():
    ordinais = ["1ª", "2ª", "3ª", "4ª", "5ª", "6ª", "7ª", "8ª", "9ª", "10ª"]
    itens = "".join(
        f"<p><strong>{ordinais[i]}</strong> - {texto}</p>"
        for i, texto in enumerate(COMPETENCIAS_GERAIS_BNCC)
    )
    return f'<h5 class="doc-secao-titulo">Competências Gerais da Educação Básica</h5><div class="doc-secao-corpo">{itens}</div>'

def sanitizar_saida_html(texto):
    """
    Rede de segurança: mesmo instruindo o Gemini a responder em HTML puro,
    modelos de linguagem eventualmente retornam sintaxe Markdown (###, **, *).
    Esta função detecta esse padrão e converte para HTML real antes de renderizar.
    """
    texto = texto.strip()

    # Heurística: se encontrarmos marcadores clássicos de Markdown, convertemos.
    parece_markdown = bool(re.search(r'(^|\n)#{2,6}\s|\*\*[^*]+\*\*|(^|\n)\*\s|(^|\n)-\s', texto))

    if parece_markdown:
        texto = md.markdown(texto, extensions=['tables', 'nl2br', 'sane_lists'])

    return texto.strip()

# =====================================================================
# SANITIZAÇÃO DE SEGURANÇA (PROTEÇÃO CONTRA XSS)
# =====================================================================
# Diferente de sanitizar_saida_html() acima (que só converte Markdown -> HTML),
# esta função REMOVE de fato tags/atributos perigosos (<script>, onerror=,
# javascript:, <iframe>, etc.) usando uma allowlist. É aplicada tanto ao HTML
# gerado pela IA quanto ao HTML digitado manualmente pelo professor na edição
# de material, antes de qualquer gravação no banco de dados.
TAGS_HTML_PERMITIDAS = [
    'h4', 'h5', 'h6', 'p', 'strong', 'em', 'b', 'i', 'u', 'br', 'hr',
    'ul', 'ol', 'li', 'div', 'span',
    'table', 'thead', 'tbody', 'tr', 'th', 'td',
    'a',
]
ATRIBUTOS_HTML_PERMITIDOS = {
    '*': ['class'],
    'a': ['href', 'title', 'target', 'rel'],
}
PROTOCOLOS_URL_PERMITIDOS = ['http', 'https', 'mailto']

def sanitizar_html_seguro(html):
    """Remove scripts, handlers de evento e outras tags/atributos perigosos,
    preservando apenas a formatação usada pelo sistema. Uso obrigatório em
    qualquer HTML que venha de fora (IA ou digitado pelo usuário) antes de
    ser salvo no banco ou renderizado com | safe."""
    if not html:
        return html
    limpo = bleach.clean(
        html,
        tags=TAGS_HTML_PERMITIDAS,
        attributes=ATRIBUTOS_HTML_PERMITIDOS,
        protocols=PROTOCOLOS_URL_PERMITIDOS,
        strip=True,
    )
    return limpo

def obter_fallback_pedagogico(tipo_modulo, tema, erro_adicional=""):
    sobrecarga = "503" in erro_adicional and ("UNAVAILABLE" in erro_adicional or "overloaded" in erro_adicional.lower() or "high demand" in erro_adicional.lower())
    limite_requisicoes = "429" in erro_adicional
    falha_conexao = "Falha de conexão física" in erro_adicional or "timed out" in erro_adicional.lower() or "timeout" in erro_adicional.lower()

    if sobrecarga:
        mensagem_principal = "Os servidores de IA estão temporariamente sobrecarregados (alta demanda no momento). O sistema já tentou os provedores configurados automaticamente, mas nenhum respondeu a tempo. Isso normalmente se resolve em poucos minutos — tente gerar novamente."
    elif limite_requisicoes:
        mensagem_principal = "Foi atingido um limite de requisições em um dos provedores de IA configurados. Aguarde um instante antes de tentar novamente, ou verifique sua cota no painel do provedor correspondente."
    elif falha_conexao:
        mensagem_principal = "O sistema não conseguiu estabelecer conexão com os servidores de IA a tempo (timeout de rede). Isso pode ser uma instabilidade temporária de rede — tente gerar novamente em alguns instantes."
    else:
        mensagem_principal = "O sistema não conseguiu gerar o conteúdo com nenhum dos provedores de IA configurados. Certifique-se de que as variáveis <strong>GEMINI_API_KEY</strong> (e, se estiver usando, <strong>MISTRAL_API_KEY</strong>) estão configuradas corretamente no painel do Render."

    return f"""
    <h4><i class="fa-solid fa-graduation-cap text-primary me-2"></i> {tipo_modulo} (Modo de Segurança)</h4>
    <p>{mensagem_principal}</p>
    <p><strong>Tema enviado:</strong> {tema}</p>
    {f'<p class="text-danger small"><strong>Detalhes do Erro:</strong> {erro_adicional}</p>' if erro_adicional else ''}
    """

def executar_geracao_ia(**kwargs):
    tipo_modulo = kwargs.get('tipo_modulo', 'Banco de Atividades')
    tema = kwargs.get('tema', '')
    disciplina = kwargs.get('disciplina', 'Geral')
    ano = kwargs.get('ano', 'Geral')
    bncc = kwargs.get('bncc', '')
    tipo_prova = kwargs.get('tipo_prova', 'Mista')
    qtd_questoes = kwargs.get('qtd_questoes', '10')
    nivel = kwargs.get('nivel', 'Médio')
    nome_professor = kwargs.get('nome_professor', 'Professor(a)')
    nome_escola = kwargs.get('nome_escola', 'Instituição de Ensino')
    usuario_email = kwargs.get('usuario_email', '')

    # Campos específicos do Planejamento Bimestral (formato oficial SEMED)
    numero_plano = kwargs.get('numero_plano', '')
    bimestre = kwargs.get('bimestre', '1º BIM')
    data_inicio = kwargs.get('data_inicio', '')
    data_fim = kwargs.get('data_fim', '')
    turma = kwargs.get('turma', '')
    turno = kwargs.get('turno', '')
    modalidade = kwargs.get('modalidade', 'Presencial')
    ano_letivo = kwargs.get('ano_letivo', '')
    inep = kwargs.get('inep', '')
    endereco_escola = kwargs.get('endereco_escola', '')
    cidade_escola = kwargs.get('cidade_escola', '')
    estado_escola = kwargs.get('estado_escola', 'MA')
    zona_escola = kwargs.get('zona_escola', '')
    telefone_escola = kwargs.get('telefone_escola', '')
    email_escola = kwargs.get('email_escola', '')
    observacoes = kwargs.get('observacoes', '')

    if not GEMINI_API_KEY:
        registrar_erro(tipo_modulo, usuario_email, "GEMINI_API_KEY ausente no painel do Render.", nivel='critico')
        return obter_fallback_pedagogico(tipo_modulo, tema, "A variável GEMINI_API_KEY está ausente no painel do Render."), '', False

    # -----------------------------------------------------------------
    # EXTRAÇÃO CIRÚRGICA DA CHAVE COM REGEX (Ignora qualquer formatação de link)
    # -----------------------------------------------------------------
    # O padrão procura pela sequência que começa com AIza ou AQ e pega apenas caracteres válidos de chave
    match = re.search(r'(AIzaSy[A-Za-z0-9_-]+|AQ\.[A-Za-z0-9_-]+)', GEMINI_API_KEY)
    
    if match:
        chave_limpa = match.group(1).strip()
    else:
        # Fallback caso a regex não isole (remove manualmente caracteres de link comuns)
        chave_limpa = GEMINI_API_KEY.replace("[", "").replace("]", "").replace("'", "").replace('"', '')
        if "key=" in chave_limpa:
            chave_limpa = chave_limpa.split("key=")[-1]
        if ")" in chave_limpa:
            chave_limpa = chave_limpa.split(")")[-1]
        chave_limpa = chave_limpa.strip()
    # -----------------------------------------------------------------

    # 2. CONSTRUÇÃO DO PROMPT PEDAGÓGICO
    regras_formato = """
        REGRAS OBRIGATÓRIAS DE FORMATAÇÃO DA SAÍDA:
        - Responda ESTRITAMENTE em HTML puro e semântico (tags: h4, h5, p, strong, em, ul, ol, li, table, thead, tbody, tr, th, td, div).
        - NUNCA utilize sintaxe Markdown (proibido: #, ##, ###, **texto**, *item*, -, ```). Se precisar de negrito use <strong>, se precisar de título use <h4>/<h5>, se precisar de lista use <ul><li>.
        - NUNCA inclua delimitadores de bloco de código como ```html ou ```.
        - Use <table class="tabela-bncc"> quando fizer sentido organizar habilidades BNCC, cronogramas ou critérios de avaliação em formato tabular.
    """
    if tipo_modulo == 'Tira-Dúvidas com IA':
        prompt = f"""
        Você é um Consultor Jurídico-Pedagógico especialista e expert em Legislação Educacional Brasileira.
        Responda com total precisão técnica fundamentando-se OBRIGATORIAMENTE em: BNCC (Base Nacional Comum Curricular), LDB (Lei nº 9.394/96), DCTMA (Documento Curricular do Território Maranhense) e, quando pertinente, na Seção da Educação da Constituição Federal.
        Sempre que citar um desses documentos, indique de forma explícita o artigo, competência ou eixo correspondente.
        Dúvida ou Consulta do Professor: "{tema}"
        {regras_formato}
        """
    elif tipo_modulo == 'Planejamento Bimestral':
        numero_final = numero_plano.strip() or str(random.randint(10000, 99999))
        ano_letivo_final = ano_letivo.strip() or str(datetime.now().year)
        criado_em = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        turma_completa = f"{ano} | ({turma}) | {turno}".upper() if turma or turno else ano.upper()
        periodo_execucao = f"{data_inicio} a {data_fim}" if data_inicio and data_fim else "[preencher período de execução]"

        prompt = f"""
        Atue como um Especialista em Planejamento Pedagógico Escolar, com domínio profundo da BNCC (Base Nacional Comum Curricular), da LDB (Lei nº 9.394/96) e do DCTMA (Documento Curricular do Território Maranhense).
        Gere um PLANEJAMENTO BIMESTRAL completo e oficial, seguindo EXATAMENTE a estrutura, ordem de seções e nível de detalhamento abaixo — este é o modelo oficial usado pela Secretaria Municipal de Educação, e deve ser reproduzido fielmente.

        {regras_formato}
        REGRA ADICIONAL CRÍTICA: Não gere você mesmo a seção "Competências Gerais da Educação Básica". No lugar exato indicado abaixo, insira apenas o marcador literal <!--COMPETENCIAS_GERAIS_AQUI--> (sem nenhum texto ao redor, sem tags H5, apenas o comentário HTML). Esse marcador será substituído automaticamente pelo sistema pelo texto oficial completo.

        ESTRUTURA OBRIGATÓRIA DO DOCUMENTO, NESTA ORDEM EXATA:

        1. <div class="doc-cabecalho-oficial">
           Inclua, em formato de linhas rotuladas (<p><strong>RÓTULO:</strong> valor</p>):
           INEP: {inep or '[preencher]'}
           ESCOLA: {nome_escola}
           ENDEREÇO: {endereco_escola or '[preencher endereço]'}
           CIDADE: {cidade_escola or '[preencher cidade]'} ESTADO: {estado_escola}
           ZONA: {zona_escola or '[preencher]'} TELEFONE: {telefone_escola or '[preencher]'} EMAIL: {email_escola or '[preencher]'}
           </div>

        2. <h2 class="doc-titulo-oficial">PLANO BIMESTRAL #{numero_final}</h2>

        3. <div class="doc-metadata-oficial">
           Em linhas rotuladas, exatamente nesta ordem:
           BIMESTRAL // {modalidade}
           {bimestre} {periodo_execucao}
           TURMA: {turma_completa}
           COMP. CURR.: {disciplina}
           EXECUÇÃO: {periodo_execucao}
           CRIADO EM: {criado_em}
           PROFESSOR(A): {nome_professor}
           ANO LETIVO: {ano_letivo_final}
           </div>

        4. <!--COMPETENCIAS_GERAIS_AQUI-->

        5. <h5 class="doc-secao-titulo">Competências Específicas de {disciplina}</h5>
           Liste, numeradas (1ª, 2ª, 3ª...), as competências específicas oficiais da BNCC para a área de conhecimento à qual "{disciplina}" pertence (ex: Ciências da Natureza, Matemática, Linguagens, Ciências Humanas), voltadas ao Ensino Fundamental — Anos Finais. Seja fiel ao texto oficial, sem inventar.

        6. <h5 class="doc-secao-titulo">Unidades Temáticas</h5>
           Identifique a(s) unidade(s) temática(s) da BNCC relacionada(s) ao tema "{tema}" para a série/ano "{ano}", no formato "Nome da Unidade / {ano}".

        7. <h5 class="doc-secao-titulo">Objetos de Conhecimento</h5>
           Liste os objetos de conhecimento relacionados ao tema, cada um seguido de "UNIDADE: [nome da unidade] / {ano}".

        8. <h5 class="doc-secao-titulo">Habilidades</h5>
           Liste as habilidades da BNCC pertinentes (formato: código - descrição completa da habilidade). Priorize o código informado pelo professor ({bncc if bncc else 'nenhum código específico informado — selecione os mais adequados ao tema'}) e complemente com habilidades relacionadas do mesmo objeto de conhecimento.

        9. <h5 class="doc-secao-titulo">Sugestões Metodológicas</h5>
           Lista <ul><li> com 8 a 10 sugestões práticas e variadas de condução das aulas ao longo do bimestre (leitura, TIC, debates, mapas conceituais, saída de campo, mostra científica, etc.), adaptadas ao tema e à realidade do Maranhão quando pertinente.

        10. <h5 class="doc-secao-titulo">Avaliação</h5>
            Lista <ul><li> com os instrumentos avaliativos do bimestre (ex: Avaliação Bimestral, Seminários, Trabalhos individuais e em grupo, etc.), adequados ao tema.

        11. <h5 class="doc-secao-titulo">Recursos</h5>
            Lista <ul><li> com os recursos didáticos necessários (materiais, equipamentos, tecnologia).

        12. <h5 class="doc-secao-titulo">Referências</h5>
            Liste em linhas simples (sem bullets), formato ABNT simplificado:
            BRASIL. Ministério da Educação. Documento Curricular do Território Maranhense (DCTMA): para o ensino fundamental. Rio de Janeiro: FGV, [ano mais recente conhecido].
            BRASIL. Ministério da Educação. Base Nacional Comum Curricular (BNCC). Brasília, 2018.
            E, se pertinente ao tema/disciplina, uma referência bibliográfica de livro didático real e amplamente reconhecido (não invente autores/obras).

        13. <h5 class="doc-secao-titulo">Observações Pertinentes</h5>
            {f'Inclua o seguinte texto informado pelo professor: "{observacoes}"' if observacoes else 'Deixe um parágrafo curto padrão indicando que não há observações adicionais registradas, ou omita se preferir deixar em branco.'}

        DADOS DE CONFIGURAÇÃO DO ESCOPO:
        - Componente/Disciplina: {disciplina}
        - Ano/Série Escolar: {ano}
        - Tema Central / Objeto de Estudo: {tema}
        - Código de Habilidade BNCC Alvo: {bncc}
        """
    elif tipo_modulo == 'Sequência Didática':
        quantidade_aulas = kwargs.get('quantidade_aulas', '5').strip() or '5'
        duracao_aula = kwargs.get('duracao', '').strip()
        objetivo_geral_informado = kwargs.get('objetivo', '').strip()
        objetivos_especificos = kwargs.get('objetivos_especificos', '').strip()
        perfil_turma = kwargs.get('perfil_turma', '').strip()
        dificuldades_turma = kwargs.get('dificuldades_turma', '').strip()
        recursos_disponiveis = kwargs.get('recursos_disponiveis', '').strip()
        metodologia_preferida = kwargs.get('metodologia_preferida', '').strip()

        try:
            qtd_aulas_num = max(1, min(int(quantidade_aulas), 20))
        except ValueError:
            qtd_aulas_num = 5

        prompt = f"""
        Atue como um Especialista em Planejamento Pedagógico, com domínio profundo da BNCC (Base Nacional Comum Curricular), da LDB (Lei nº 9.394/96) e do DCTMA (Documento Curricular do Território Maranhense).
        Gere uma SEQUÊNCIA DIDÁTICA completa sobre o tema "{tema}", disciplina "{disciplina}", ano/série "{ano}", com exatamente {qtd_aulas_num} aula(s).

        {regras_formato}
        REGRA ADICIONAL CRÍTICA: Não gere você mesmo a seção "Competências Gerais da Educação Básica". No local exato indicado abaixo, insira apenas o marcador literal <!--COMPETENCIAS_GERAIS_AQUI--> (sem nenhum texto ao redor, sem tags H5, apenas o comentário HTML). Esse marcador será substituído automaticamente pelo sistema pelo texto oficial completo.

        ESTRUTURA OBRIGATÓRIA DO DOCUMENTO, NESTA ORDEM EXATA, cada uma como <h5 class="doc-secao-titulo">[Título]</h5> seguida do conteúdo:

        1. Identificação — em linhas rotuladas (<p><strong>RÓTULO:</strong> valor</p>): Disciplina: {disciplina}; Ano/Série: {ano}; Duração total: {qtd_aulas_num} aula(s){f', {duracao_aula} cada' if duracao_aula else ''}; Professor(a): {nome_professor}; Escola: {nome_escola}.
        2. Tema — apresente o tema "{tema}" de forma clara, contextualizando sua relevância para o ano/série informado.
        3. Justificativa — um parágrafo explicando por que este tema é importante para os alunos desta etapa, com base pedagógica.
        4. Objetivo Geral — {f'considere o objetivo informado pelo professor: "{objetivo_geral_informado}"' if objetivo_geral_informado else 'formule um objetivo geral claro e mensurável para a sequência'}.
        5. Objetivos Específicos — lista <ul><li> com 3 a 6 objetivos específicos. {f'Considere os informados pelo professor: "{objetivos_especificos}"' if objetivos_especificos else 'Elabore objetivos coerentes com o objetivo geral e o tema.'}
        6. Habilidades BNCC — liste as habilidades da BNCC pertinentes (formato: código - descrição completa). Priorize o código informado pelo professor ({bncc if bncc else 'nenhum código específico informado — selecione os mais adequados ao tema e ano/série'}).
        7. Conteúdos — lista <ul><li> com os conteúdos conceituais, procedimentais e atitudinais trabalhados ao longo da sequência.
        8. Metodologia — descreva a abordagem metodológica geral da sequência. {f'O professor indicou preferência por: "{metodologia_preferida}"' if metodologia_preferida else 'Escolha uma abordagem ativa e adequada ao tema (ex: sala de aula invertida, aprendizagem baseada em problemas, rotação por estações).'}
        9. Desenvolvimento Aula por Aula — use <div class="questao-item"> para cada aula (uma div por aula). Gere exatamente {qtd_aulas_num} blocos, cada um no formato "<strong>Aula 1 — [nome curto da etapa, ex: Introdução]</strong>" seguido de um parágrafo descrevendo o que será feito naquela aula (atividades, tempo estimado, interação com os alunos). Distribua a progressão pedagógica de forma lógica ao longo das aulas (ex: introdução → desenvolvimento → prática → sistematização → avaliação, adaptando à quantidade de aulas informada).
        10. Recursos — lista <ul><li> com os recursos didáticos necessários. {f'O professor tem disponível: "{recursos_disponiveis}"' if recursos_disponiveis else 'Sugira recursos comuns e acessíveis à realidade escolar pública.'}
        11. Avaliação — lista <ul><li> com os instrumentos e critérios de avaliação da aprendizagem ao longo da sequência.
        12. Inclusão/Adaptações — sugestões de adaptação para alunos com dificuldades de aprendizagem ou necessidades específicas, SEM realizar qualquer diagnóstico médico ou clínico — trabalhe apenas com orientações pedagógicas gerais. {f'Perfil da turma informado pelo professor: "{perfil_turma}".' if perfil_turma else ''} {f'Dificuldades observadas: "{dificuldades_turma}".' if dificuldades_turma else ''}
        13. Referências — em linhas simples (sem bullets), formato ABNT simplificado, citando BNCC (2018), a LDB (Lei nº 9.394/96) e, se pertinente, o DCTMA.
        """
    elif tipo_modulo == 'Gerador de Provas':
        prompt = f"""
        Atue como um Especialista em Avaliação Pedagógica, com domínio profundo da BNCC, da LDB (Lei nº 9.394/96) e do DCTMA (Documento Curricular do Território Maranhense).
        Gere uma PROVA/AVALIAÇÃO completa sobre o tema "{tema}", disciplina "{disciplina}", ano/série "{ano}".
        NÃO gere nenhum cabeçalho com nome de escola, professor, aluno ou data — isso já é adicionado automaticamente pelo sistema.

        {regras_formato}

        SUA RESPOSTA DEVE TER DUAS PARTES, NESTA ORDEM, SEPARADAS PELO MARCADOR LITERAL EXATO <!--INFO_PEDAGOGICA--> (sem nenhum texto ao redor do marcador):

        ============ PARTE 1 — A PROVA EM SI (vai para o documento impresso/exportado) ============
        1. Gere exatamente {qtd_questoes} questões no formato de aplicação: {tipo_prova}.
        2. Cada questão DEVE estar envolvida em <div class="questao-item">...</div> (uma div por questão, isso é obrigatório para a diagramação em colunas).
        3. Utilize estritamente numeração sequencial de dois dígitos seguida de ponto (Exemplo: 01., 02., 03.).
        4. Sempre inclua a diretriz BNCC entre parênteses logo após o número. Exemplo: '01. (EF09MA02) '.
        5. Todo o texto do enunciado da pergunta DEVE estar encapsulado dentro da tag HTML <strong>...</strong>.
        6. Para questões objetivas, organize alternativas perfeitamente alinhadas verticalmente de a) até d) separadas por quebras de linha <br>.
        7. Para questões discursivas ou subjetivas, estime mentalmente quantas linhas a resposta esperada ocuparia (com base na complexidade da pergunta) e gere um número de <div class="linha-resposta"></div> igual a esse número + 3 linhas extras (exemplo: se a resposta esperada ocupa cerca de 5 linhas, gere 8 divs). Nunca gere menos de 5 linhas no total, para garantir espaço confortável ao aluno.
        8. Após a última questão, inclua <div class="gabarito-prova"><h5>Gabarito</h5>[resposta correta de cada questão, apenas número + alternativa, de forma resumida]</div>.

        ============ PARTE 2 — INFORMAÇÕES PEDAGÓGICAS (fica visível só na tela do sistema, NUNCA é exportada/impressa) ============
        Após o marcador <!--INFO_PEDAGOGICA-->, gere, cada uma como <h5>[Título]</h5>:
        - Objetivos: o que se espera avaliar com esta prova.
        - Competências e Habilidades da BNCC: competência(s) geral(is) e habilidade(s) específica(s) (código + descrição) trabalhadas, priorizando o código informado ({bncc if bncc else 'nenhum informado — selecione os mais adequados ao tema'}).
        - Fundamentação Legal: artigo pertinente da LDB (Lei 9.394/96) e eixo/orientação do DCTMA relacionados ao tema.
        - Configuração Utilizada: resuma em lista os parâmetros desta prova — Disciplina: {disciplina}; Ano/Série: {ano}; Nível: {nivel}; Formato: {tipo_prova}; Quantidade de questões: {qtd_questoes}.
        - Metodologia de Aplicação: sugestões de como aplicar esta avaliação em sala (tempo sugerido, orientações antes da aplicação, critérios de correção).
        """
    elif tipo_modulo == 'Simulados':
        tipo_simulado = kwargs.get('tipo_simulado', 'Simulado Geral / Diagnóstico')
        duracao_simulado = kwargs.get('duracao_simulado', '').strip()
        qtd_questoes_simulado = kwargs.get('qtd_questoes_simulado', '20').strip() or '20'
        duracao_sugerida = duracao_simulado
        if not duracao_sugerida:
            try:
                duracao_sugerida = f"aproximadamente {int(qtd_questoes_simulado) * 3} minutos"
            except ValueError:
                duracao_sugerida = "a critério do professor"

        prompt = f"""
        Atue como um Especialista em Avaliação Pedagógica e Elaboração de Simulados, com domínio profundo da BNCC, da LDB (Lei nº 9.394/96) e do DCTMA (Documento Curricular do Território Maranhense).
        Gere um SIMULADO completo sobre o(s) tema(s)/conteúdo(s) "{tema}", disciplina "{disciplina}", ano/série "{ano}".
        Tipo de simulado: {tipo_simulado}.
        NÃO gere nenhum cabeçalho com nome de escola, professor, aluno ou data — isso já é adicionado automaticamente pelo sistema.

        {regras_formato}

        SUA RESPOSTA DEVE TER DUAS PARTES, NESTA ORDEM, SEPARADAS PELO MARCADOR LITERAL EXATO <!--INFO_PEDAGOGICA--> (sem nenhum texto ao redor do marcador):

        ============ PARTE 1 — O SIMULADO EM SI (vai para o documento impresso/exportado) ============
        1. Inicie com <div class="instrucoes-simulado"><p><strong>Duração sugerida:</strong> {duracao_sugerida}</p><p><strong>Instruções:</strong> [orientações curtas e claras ao aluno sobre como responder, adequadas ao tipo de simulado "{tipo_simulado}"]</p></div>
        2. Gere exatamente {qtd_questoes_simulado} questões, cobrindo de forma equilibrada o(s) conteúdo(s)/tema(s) informado(s). Se mais de um conteúdo/tema foi informado, distribua as questões proporcionalmente entre eles.
        3. Varie o nível de dificuldade ao longo do simulado (fácil, médio, difícil) de forma progressiva ou intercalada, mesmo que o nível informado seja único — isso é característico de simulados reais. Se o nível informado for "Misto", isso é ainda mais importante.
        4. Cada questão DEVE estar envolvida em <div class="questao-item">...</div> (uma div por questão, obrigatório para a diagramação em colunas).
        5. Utilize estritamente numeração sequencial de dois dígitos seguida de ponto (Exemplo: 01., 02., 03.).
        6. Sempre inclua a diretriz BNCC entre parênteses logo após o número. Exemplo: '01. (EF09MA02) '.
        7. Todo o texto do enunciado da pergunta DEVE estar encapsulado dentro da tag HTML <strong>...</strong>.
        8. Para questões objetivas, organize alternativas perfeitamente alinhadas verticalmente de a) até d) separadas por quebras de linha <br>.
        9. Para questões discursivas, estime mentalmente quantas linhas a resposta esperada ocuparia (com base na complexidade da pergunta) e gere um número de <div class="linha-resposta"></div> igual a esse número + 3 linhas extras (exemplo: se a resposta esperada ocupa cerca de 5 linhas, gere 8 divs). Nunca gere menos de 5 linhas no total, para garantir espaço confortável ao aluno.
        10. Após a última questão, inclua <div class="gabarito-prova"><h5>Gabarito</h5>[resposta correta de cada questão, apenas número + alternativa, de forma resumida]</div>.

        ============ PARTE 2 — INFORMAÇÕES PEDAGÓGICAS (fica visível só na tela do sistema, NUNCA é exportada/impressa) ============
        Após o marcador <!--INFO_PEDAGOGICA-->, gere, cada uma como <h5>[Título]</h5>:
        - Objetivos: o que se espera diagnosticar/revisar com este simulado.
        - Matriz de Referência: uma <table> com colunas "Questão", "Habilidade BNCC", "Nível de Dificuldade", uma linha por questão gerada — replicando o formato de matriz de referência usado em simulados oficiais (tipo SAEB/ENEM).
        - Fundamentação Legal: artigo pertinente da LDB (Lei 9.394/96) e eixo/orientação do DCTMA relacionados ao(s) tema(s).
        - Configuração Utilizada: resuma em lista — Disciplina: {disciplina}; Ano/Série: {ano}; Tipo de Simulado: {tipo_simulado}; Nível: {nivel}; Quantidade de questões: {qtd_questoes_simulado}; Duração sugerida: {duracao_sugerida}.
        - Metodologia de Aplicação: sugestões de como aplicar este simulado (condições de sala, correção, uso dos resultados para intervenção pedagógica).
        """
    else:
        prompt = f"""
        Atue como um Especialista em Design Pedagógico e Elaboração de Conteúdo Escolar Avançado, com domínio profundo da BNCC, da LDB (Lei nº 9.394/96) e do DCTMA (Documento Curricular do Território Maranhense).
        Gere o conteúdo completo e detalhado para o documento estruturado do módulo '{tipo_modulo}'.
        NÃO gere nenhum cabeçalho com nome de escola, professor, aluno ou data — o cabeçalho do documento já é adicionado automaticamente pelo sistema. Comece diretamente pelo conteúdo pedagógico.

        FUNDAMENTAÇÃO LEGAL OBRIGATÓRIA:
        Todo o conteúdo pedagógico deve estar alinhado e referenciar explicitamente, quando aplicável:
        1. BNCC — cite a(s) competência(s) geral(is) e a(s) habilidade(s) específica(s) trabalhada(s).
        2. LDB (Lei nº 9.394/96) — cite o artigo pertinente aos princípios/fins da educação relacionados ao tema.
        3. DCTMA — cite o eixo ou orientação curricular do território maranhense relacionado.
        Inclua uma seção final <h5>Fundamentação Legal</h5> resumindo essas referências.

        DADOS DE CONFIGURAÇÃO DO ESCOPO:
        - Componente/Disciplina: {disciplina}
        - Ano/Série Escolar: {ano}
        - Tema Central / Objeto de Estudo: {tema}
        - Código de Habilidade BNCC Alvo: {bncc}
        - Nível de Rigor Cognitivo: {nivel}
        {regras_formato}
        """
        if tipo_modulo == 'Banco de Atividades':
            tipos_selecionados = kwargs.get('tipos_atividade', []) or ['Questões objetivas']
            lista_tipos = ", ".join(tipos_selecionados)
            prompt += f"""
            DIRETRIZES DO BANCO DE ATIVIDADES:
            1. O professor selecionou os seguintes TIPOS DE ATIVIDADE, que devem OBRIGATORIAMENTE estar presentes no material, cada um em sua própria seção com <h5>[Nome do tipo de atividade]</h5>: {lista_tipos}.
            2. Adapte cada tipo de atividade ao ano/série "{ano}" e ao tema "{tema}" — para anos iniciais/alfabetização, use atividades mais visuais e lúdicas; para anos finais/médio, mais analíticas.
            3. Se o nível de dificuldade for "Misto", varie a dificuldade das questões dentro de cada seção (inclua fácil, médio e difícil).
            4. Para "Complete as lacunas", use "_______" para os espaços em branco.
            5. Para "Ligue as colunas", monte duas colunas usando uma <table>.
            6. Para "Verdadeiro ou Falso", numere as afirmações e inclua parênteses "( )" antes de cada uma para marcação.
            7. Para "Caça-palavras" e "Cruzadinha", gere a LISTA de palavras-chave relacionadas ao tema (não é necessário desenhar a grade visualmente, apenas a lista de palavras e uma breve instrução de como montá-la).
            8. Para atividades práticas/ilustradas/recorte e cole/pintura educativa, descreva claramente as instruções passo a passo para o professor aplicar em sala.
            9. Ao final de TODAS as seções com questões que tenham resposta certa (objetivas, lacunas, V/F, ligue as colunas, sequência lógica, problemas matemáticos), inclua uma seção final única <h5>Gabarito</h5> reunindo as respostas de todas essas atividades.
            """
        elif tipo_modulo == 'Projetos Interdisciplinares':
            duracao = kwargs.get('duracao', '')
            objetivo_geral_prof = kwargs.get('objetivo', '')
            tipos_projeto_sel = kwargs.get('tipos_projeto', []) or ['Projeto Interdisciplinar']
            lista_tipos_projeto = ", ".join(tipos_projeto_sel)
            prompt += f"""
            DIRETRIZES DO GERADOR DE PROJETOS PEDAGÓGICOS:
            O professor quer um projeto do(s) seguinte(s) tipo(s)/formato(s): {lista_tipos_projeto}. Combine as características desses formatos de forma coerente se mais de um for selecionado.
            Duração prevista informada pelo professor: {duracao or 'não informada — sugira uma duração adequada ao escopo do tema'}.
            Objetivo inicial informado pelo professor (use como norte, mas expanda): {objetivo_geral_prof or 'não informado — infira a partir do tema'}.

            ESTRUTURA OBRIGATÓRIA, NESTA ORDEM, cada uma como <h5>[Título da Seção]</h5>:
            1. Justificativa
            2. Objetivo Geral
            3. Objetivos Específicos (lista <ul><li>)
            4. Competências da BNCC (numeradas, relacionadas ao tema e à área)
            5. Habilidades da BNCC (código + descrição, relacionadas ao tema e ano/série "{ano}")
            6. Metodologia
            7. Cronograma (apresente como <table> com colunas Etapa / Descrição / Período, distribuído ao longo da duração informada)
            8. Recursos (lista <ul><li>)
            9. Desenvolvimento (descrição detalhada de como o projeto se desenrola do início ao fim)
            10. Avaliação
            11. Produto Final (o que será entregue/apresentado ao final)
            12. Referências (formato ABNT simplificado, incluindo BNCC e DCTMA; não invente autores)
            """
        elif tipo_modulo == 'Alfabetização e Reforço de Leitura':
            nome_aluno = kwargs.get('nome_aluno', '').strip()
            nivel_leitura = kwargs.get('nivel_leitura', '')
            dificuldades_observadas = kwargs.get('dificuldades_observadas', '')
            focos_selecionados = kwargs.get('focos_alfabetizacao', []) or ['Consciência Fonológica (rimas, sons, sílabas)']
            lista_focos = ", ".join(focos_selecionados)
            duracao_alfab = kwargs.get('duracao_alfabetizacao', '')
            eh_aluno_mais_velho = False
            try:
                eh_aluno_mais_velho = int(re.search(r'\d+', ano).group()) >= 4 if ano and re.search(r'\d+', ano) else False
            except Exception:
                eh_aluno_mais_velho = False

            prompt += f"""
            DIRETRIZES DO PLANO DE ALFABETIZAÇÃO E REFORÇO DE LEITURA:
            Este é um plano de INTERVENÇÃO INDIVIDUALIZADA para um aluno com defasagem de leitura/escrita, não uma atividade de sala genérica.

            DADOS DO ALUNO:
            - Nome (se informado): {nome_aluno or 'não identificado — trate de forma genérica como "o(a) aluno(a)"'}
            - Ano/Série atual: {ano}
            - Nível de leitura/escrita diagnosticado pelo professor: {nivel_leitura or 'não informado — infira um nível plausível a partir das dificuldades descritas'}
            - Dificuldades observadas pelo professor: {dificuldades_observadas or 'não detalhadas — baseie-se no nível informado'}
            - Foco(s) de intervenção priorizado(s): {lista_focos}
            - Duração prevista do plano: {duracao_alfab or 'sugira uma duração adequada (geralmente 4 a 8 semanas)'}
            - Contexto/interesses do aluno informados pelo professor (use para escolher temas de textos e exemplos motivadores): {tema or 'não informado — escolha temas neutros e universalmente interessantes para a idade'}

            {"ATENÇÃO CRÍTICA: o aluno está no " + ano + ", ou seja, é mais velho que a idade típica de alfabetização inicial. NÃO use temas, personagens ou materiais infantilizados (nada de 'bichinhos fofos' ou temas de educação infantil). Use textos, palavras e contextos adequados à idade real do aluno (esportes, música, tecnologia, cotidiano adolescente, temas de interesse da faixa etária), mesmo trabalhando habilidades básicas de leitura. Isso é essencial para não constranger o aluno diante dos colegas." if eh_aluno_mais_velho else "Adeque a linguagem e os temas à faixa etária da educação infantil/anos iniciais, com abordagem lúdica."}

            ESTRUTURA OBRIGATÓRIA, cada uma como <h5>[Título]</h5>:
            1. Diagnóstico e Leitura da Situação — interprete o nível/dificuldades informados e explique o que isso significa na prática.
            2. Objetivos da Intervenção (geral e específicos).
            3. Habilidades da BNCC Relacionadas — competências/habilidades de Língua Portuguesa (alfabetização/leitura) pertinentes ao nível diagnosticado, com código quando aplicável.
            4. Estratégias e Metodologia — métodos de alfabetização reconhecidos (consciência fonológica, método fônico, silabação, leitura compartilhada, etc.), adaptados ao nível e à idade real do aluno.
            5. Sequência de Atividades Práticas — organize em blocos (ex: Semana 1, Semana 2...) com atividades concretas e progressivas, do mais simples ao mais complexo.
            6. Recursos Necessários.
            7. Envolvimento da Família — orientações simples para os responsáveis reforçarem em casa.
            8. Avaliação de Progresso — como o professor vai medir a evolução (indicadores observáveis, não apenas provas).
            9. Referências (formato ABNT simplificado; inclua BNCC, e se pertinente, autores reconhecidos de alfabetização como Magda Soares ou Emilia Ferreiro — não invente nomes/obras).
            """
        elif tipo_modulo == 'Plano de Aula':
            prompt += """
            SEÇÃO OBRIGATÓRIA — METODOLOGIAS E SUGESTÕES DE AULAS DIFERENCIADAS:
            Inclua, antes da conclusão do plano, uma seção <h5>Sugestões de Aulas Diferenciadas</h5> com pelo menos 3 a 4 propostas concretas e variadas para trabalhar o tema de formas alternativas à aula expositiva tradicional, por exemplo (adapte ao tema e ano/série informados):
            - Aula prática/experimental (uso de materiais concretos, experimentos, manipuláveis).
            - Metodologia ativa (sala de aula invertida, aprendizagem baseada em problemas/projetos, gamificação).
            - Atividade em grupo/colaborativa (debate, júri simulado, oficina, estudo de caso).
            - Uso de tecnologia/recursos digitais (aplicativos, vídeos, simuladores, jogos educativos).
            - Conexão com o cotidiano/comunidade local (saída de campo, entrevista, estudo do meio, quando aplicável ao contexto maranhense).
            Apresente cada sugestão em formato de lista <ul><li>, com um parágrafo curto (2-3 linhas) explicando como aplicá-la e qual habilidade/competência da BNCC ela reforça.

            SEÇÃO FINAL OBRIGATÓRIA — REFERÊNCIAS:
            Ao final do documento, após todo o conteúdo pedagógico, inclua uma seção <h5>Referências</h5> contendo:
            1. As referências normativas/legais utilizadas, no formato ABNC simplificado, por exemplo:
               BRASIL. Ministério da Educação. Base Nacional Comum Curricular (BNCC). Brasília: MEC, 2018.
               BRASIL. Lei nº 9.394, de 20 de dezembro de 1996. Lei de Diretrizes e Bases da Educação Nacional (LDB). Brasília, 1996.
               MARANHÃO. Secretaria de Estado da Educação. Documento Curricular do Território Maranhense (DCTMA). São Luís, [ano de publicação mais recente conhecido].
            2. Quaisquer referências bibliográficas pedagógicas adicionais (autores, teóricos ou materiais didáticos) efetivamente utilizados como base conceitual para o conteúdo gerado, também em formato ABNT simplificado, listadas em <ul><li>.
            3. Não invente nomes de autores ou obras específicas que não sejam amplamente reconhecidas na área; se não houver referência bibliográfica adicional além das normativas, inclua apenas as normativas.
            """
        else:
            prompt += f"\nEstruture o documento de forma oficial e profissional com cabeçalhos h4, h5, parágrafos bem espaçados e listas dinâmicas."

    # 3. CADEIA DE PROVEDORES DE IA — fallback automático entre modelos e serviços.
    # Se o Gemini (principal) estiver sobrecarregado (503) ou no limite (429) após
    # as tentativas, o sistema cai para o próximo da lista — incluindo, por último,
    # o Mistral como reforço gratuito extra — antes de exibir o Modo de Segurança.
    payload_gemini = {"contents": [{"parts": [{"text": prompt}]}]}
    headers_gemini = {'Content-Type': 'application/json'}

    mistral_chave_limpa = MISTRAL_API_KEY.replace("[", "").replace("]", "").replace("'", "").replace('"', "").strip()
    payload_mistral = {
        "model": "mistral-small-latest",
        "messages": [{"role": "user", "content": prompt}]
    }
    headers_mistral = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {mistral_chave_limpa}'
    }

    def _extrair_texto_gemini(resultado):
        return resultado['candidates'][0]['content']['parts'][0]['text']

    def _extrair_texto_mistral(resultado):
        return resultado['choices'][0]['message']['content']

    PROVEDORES_EM_CASCATA = [
        {
            "nome": "gemini-3.5-flash",
            "url": f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={chave_limpa}",
            "headers": headers_gemini, "payload": payload_gemini,
            "extrair": _extrair_texto_gemini, "ativo": bool(chave_limpa),
        },
        {
            "nome": "gemini-2.5-flash-lite",
            "url": f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={chave_limpa}",
            "headers": headers_gemini, "payload": payload_gemini,
            "extrair": _extrair_texto_gemini, "ativo": bool(chave_limpa),
        },
        {
            "nome": "mistral-small-latest",
            "url": "https://api.mistral.ai/v1/chat/completions",
            "headers": headers_mistral, "payload": payload_mistral,
            "extrair": _extrair_texto_mistral, "ativo": bool(mistral_chave_limpa),
        },
    ]

    MAX_TENTATIVAS_POR_PROVEDOR = 2
    ultimo_erro = ""

    for provedor in PROVEDORES_EM_CASCATA:
        if not provedor["ativo"]:
            continue  # chave não configurada para este provedor -> pula direto para o próximo

        for tentativa in range(1, MAX_TENTATIVAS_POR_PROVEDOR + 1):
            try:
                response = requests.post(
                    provedor["url"], headers=provedor["headers"], json=provedor["payload"], timeout=60
                )

                if response.status_code == 200:
                    resultado = response.json()
                    texto_gerado = provedor["extrair"](resultado)
                    texto_gerado = texto_gerado.replace("```html", "").replace("```", "").strip()
                    texto_gerado = sanitizar_saida_html(texto_gerado)
                    texto_gerado = sanitizar_html_seguro(texto_gerado)
                    # Substitui o marcador pelo texto oficial fixo das Competências Gerais da Educação Básica
                    if '<!--COMPETENCIAS_GERAIS_AQUI-->' in texto_gerado:
                        texto_gerado = texto_gerado.replace('<!--COMPETENCIAS_GERAIS_AQUI-->', montar_html_competencias_gerais())

                    info_pedagogica = ''
                    if tipo_modulo in ('Gerador de Provas', 'Simulados'):
                        if '<!--INFO_PEDAGOGICA-->' in texto_gerado:
                            parte_prova, parte_info = texto_gerado.split('<!--INFO_PEDAGOGICA-->', 1)
                        else:
                            parte_prova, parte_info = texto_gerado, ''
                        texto_gerado, gabarito_extraido = montar_prova_duas_colunas(parte_prova)
                        bloco_gabarito_tela = f'<div class="gabarito-tela">{gabarito_extraido}</div>' if gabarito_extraido else ''
                        info_pedagogica = (bloco_gabarito_tela + parte_info.strip()).strip()

                    return texto_gerado, info_pedagogica, True

                # Erros temporários: sobrecarga (503) ou limite de requisições (429).
                # Tenta de novo no mesmo provedor com backoff progressivo antes de trocar.
                if response.status_code in (503, 429):
                    ultimo_erro = f"Código {response.status_code} - Provedor {provedor['nome']} - Resposta: {response.text}"
                    if tentativa < MAX_TENTATIVAS_POR_PROVEDOR:
                        time.sleep(2 * tentativa)
                        continue
                    break  # esgotou tentativas neste provedor -> tenta o próximo da cascata

                # Erro definitivo (ex.: 400, 404) — não adianta insistir neste provedor,
                # mas ainda vale tentar o próximo da cascata, caso exista.
                ultimo_erro = f"Código {response.status_code} - Provedor {provedor['nome']} - Resposta: {response.text}"
                break

            except Exception as e:
                ultimo_erro = f"Falha de conexão física - Provedor {provedor['nome']}: {str(e)}"
                if tentativa < MAX_TENTATIVAS_POR_PROVEDOR:
                    time.sleep(2 * tentativa)
                    continue
                break  # esgotou tentativas neste provedor -> tenta o próximo da cascata

    registrar_erro(tipo_modulo, usuario_email, ultimo_erro + " (cascata de provedores esgotada)", nivel='atencao')
    return obter_fallback_pedagogico(tipo_modulo, tema, ultimo_erro + " (cascata de provedores esgotada)"), '', False

# =====================================================================
# EXPORTAÇÃO — DOCX e PDF (preservando cabeçalhos, tabelas, listas, negrito)
# =====================================================================
def _definir_colunas_secao(section, num_colunas):
    """Configura o número de colunas (estilo jornal) de uma seção do Word via XML."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_colunas))
    cols.set(qn('w:space'), '480')

def _adicionar_no_docx(documento, node):
    """Converte um nó de questão (strong/br/div.linha-resposta/texto) em um parágrafo do Word."""
    p = documento.add_paragraph()
    for filho in node.children:
        nome = getattr(filho, 'name', None)
        if nome == 'strong' or nome == 'b':
            run = p.add_run(filho.get_text())
            run.bold = True
        elif nome == 'br':
            p.add_run().add_break()
        elif nome == 'div' and 'linha-resposta' in (filho.get('class') or []):
            documento.add_paragraph('_' * 55)
        elif nome is None:
            texto = str(filho)
            if texto.strip():
                p.add_run(texto)
        else:
            texto = filho.get_text()
            if texto.strip():
                p.add_run(texto)
    documento.add_paragraph("")

def gerar_docx(titulo, escola, professor, html_conteudo, tipo_modulo='', disciplina='', ano=''):
    documento = Document()

    p_escola = documento.add_paragraph()
    p_escola.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_escola = p_escola.add_run(escola or "Instituição de Ensino")
    r_escola.bold = True
    r_escola.font.size = Pt(15)

    if tipo_modulo in ('Gerador de Provas', 'Simulados'):
        rotulo_titulo = "AVALIAÇÃO DE" if tipo_modulo == 'Gerador de Provas' else "SIMULADO DE"
        p_titulo = documento.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_titulo = p_titulo.add_run(f"{rotulo_titulo} {(disciplina or '').upper()}" + (f" — {ano}" if ano else ""))
        r_titulo.bold = True
        r_titulo.font.size = Pt(13)

        p_prof = documento.add_paragraph()
        p_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_prof = p_prof.add_run(f"Professor(a): {professor}")
        r_prof.italic = True
        r_prof.font.size = Pt(10)

        documento.add_paragraph("")

        # Campos de identificação (Aluno / Turma / Data / Nota)
        tabela_id = documento.add_table(rows=2, cols=2)
        tabela_id.cell(0, 0).text = "ALUNO(A): " + "_" * 38
        tabela_id.cell(0, 1).text = "TURMA: " + "_" * 10
        tabela_id.cell(1, 0).text = "DATA: ____/____/______"
        tabela_id.cell(1, 1).text = "NOTA: " + "_" * 12
        documento.add_paragraph("")

        soup = BeautifulSoup(html_conteudo, 'html.parser')
        itens = soup.find_all('div', class_='questao-item')
        gabarito_tag = soup.find('div', class_='gabarito-prova')

        # Seção em 2 colunas reais do Word para as questões
        secao_questoes = documento.add_section(WD_SECTION.CONTINUOUS)
        _definir_colunas_secao(secao_questoes, 2)

        estilo_normal = documento.styles['Normal']
        estilo_normal.font.name = 'Arial'
        estilo_normal.font.size = Pt(12)

        if itens:
            for item in itens:
                _adicionar_no_docx(documento, item)
        else:
            documento.add_paragraph(soup.get_text())

        if gabarito_tag:
            # Volta para 1 coluna antes do gabarito
            secao_gabarito = documento.add_section(WD_SECTION.CONTINUOUS)
            _definir_colunas_secao(secao_gabarito, 1)

            documento.add_paragraph("")
            h_gab = documento.add_paragraph()
            r_gab = h_gab.add_run("Gabarito")
            r_gab.bold = True
            r_gab.font.size = Pt(12)
            gabarito_copia = BeautifulSoup(str(gabarito_tag), 'html.parser')
            titulo_existente = gabarito_copia.find('h5')
            if titulo_existente:
                titulo_existente.decompose()
            conversor = HtmlToDocx()
            conversor.add_html_to_document(str(gabarito_copia), documento)

    else:
        p_titulo = documento.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_titulo = p_titulo.add_run(titulo or tipo_modulo or "Documento Pedagógico")
        r_titulo.bold = True
        r_titulo.font.size = Pt(14)

        p_prof = documento.add_paragraph()
        p_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_prof = p_prof.add_run(f"Professor(a): {professor}")
        r_prof.italic = True

        documento.add_paragraph("")
        conversor = HtmlToDocx()
        conversor.add_html_to_document(html_conteudo, documento)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_pdf(titulo, escola, professor, html_conteudo, tipo_modulo='', disciplina='', ano=''):
    if tipo_modulo in ('Gerador de Provas', 'Simulados'):
        rotulo_titulo_pdf = "Avaliação de" if tipo_modulo == 'Gerador de Provas' else "Simulado de"
        bloco_cabecalho = f"""
        <div class="cabecalho-pdf">
            <h1>{escola}</h1>
            <h2 style="font-size:13pt; text-transform:uppercase; letter-spacing:1px; margin:4px 0;">{rotulo_titulo_pdf} {(disciplina or '').upper()}{f' — {ano}' if ano else ''}</h2>
            <p style="font-size:10pt; color:#52627e;">Professor(a): {professor}</p>
            <table style="border:none; margin-top:10px;">
                <tr>
                    <td style="border:none; width:65%;">ALUNO(A): {'_' * 45}</td>
                    <td style="border:none; width:35%;">TURMA: {'_' * 10}</td>
                </tr>
                <tr>
                    <td style="border:none;">DATA: ____/____/______</td>
                    <td style="border:none;">NOTA: {'_' * 12}</td>
                </tr>
            </table>
        </div>
        """
    else:
        bloco_cabecalho = f"""
        <div class="cabecalho-pdf">
            <h1>{titulo or 'Documento Pedagógico'}</h1>
            <p><strong>Instituição de Ensino:</strong> {escola}<br/><strong>Professor(a):</strong> {professor}</p>
        </div>
        """

    html_completo = f"""
    <html>
    <head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; color: #0f1f3d; }}
        h1 {{ font-size: 16pt; text-align: center; color: #0e2a5e; margin-bottom: 2px; }}
        h4, h5 {{ color: #123a7a; margin-top: 14px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
        th, td {{ border: 1px solid #999; padding: 5px 8px; font-size: 10pt; }}
        th {{ background-color: #eef2fa; }}
        .cabecalho-pdf {{ text-align: center; margin-bottom: 18px; border-bottom: 1px solid #999; padding-bottom: 10px; }}
    </style>
    </head>
    <body>
        {bloco_cabecalho}
        {html_conteudo}
    </body>
    </html>
    """
    buffer = BytesIO()
    pisa.CreatePDF(src=html_completo, dest=buffer, encoding='utf-8')
    buffer.seek(0)
    return buffer

# =====================================================================
# INICIALIZAÇÃO DO BANCO DE DADOS (SQLite)
# =====================================================================
def init_db():
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            escola TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL
        )
    ''')
    cursor.execute("SELECT * FROM usuarios WHERE LOWER(email) = 'samuel.ssousa1506@gmail.com'")
    if not cursor.fetchone():
        cursor.execute('''
            INSERT INTO usuarios (nome, escola, email, senha) 
            VALUES ('Samuel Araújo Sousa', 'U.E. Prof. João Martins Neto', 'samuel.ssousa1506@gmail.com', '123456')
        ''')

    # Histórico + Biblioteca (favoritos) de materiais gerados
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS materiais (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_email TEXT NOT NULL,
            tipo_modulo TEXT,
            titulo TEXT,
            disciplina TEXT,
            ano TEXT,
            conteudo_html TEXT,
            favorito INTEGER DEFAULT 0,
            criado_em TEXT
        )
    ''')

    # Banco de Questões reutilizáveis
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS banco_questoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_email TEXT NOT NULL,
            ano TEXT,
            disciplina TEXT,
            conteudo TEXT,
            bncc TEXT,
            dificuldade TEXT,
            enunciado_html TEXT,
            criado_em TEXT
        )
    ''')

    # Migração segura: adiciona a coluna info_pedagogica se o banco já existia sem ela
    try:
        cursor.execute("ALTER TABLE materiais ADD COLUMN info_pedagogica TEXT")
    except sqlite3.OperationalError:
        pass  # coluna já existe

    # -----------------------------------------------------------------
    # MIGRAÇÃO v2.0 — BIBLIOTECA 2.0 (pastas)
    # -----------------------------------------------------------------
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS pastas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_email TEXT NOT NULL,
            nome TEXT NOT NULL,
            criado_em TEXT
        )
    ''')
    try:
        cursor.execute("ALTER TABLE materiais ADD COLUMN pasta_id INTEGER")
    except sqlite3.OperationalError:
        pass  # coluna já existe

    # -----------------------------------------------------------------
    # MIGRAÇÃO v2.0 — FUNDAÇÃO (aditiva, nunca apaga dados existentes)
    # -----------------------------------------------------------------
    # 1. Novas colunas da tabela usuarios: papel de acesso, status da conta
    #    e timestamps de auditoria.
    migracoes_usuarios = [
        "ALTER TABLE usuarios ADD COLUMN role TEXT NOT NULL DEFAULT 'professor'",
        "ALTER TABLE usuarios ADD COLUMN ativo INTEGER NOT NULL DEFAULT 1",
        "ALTER TABLE usuarios ADD COLUMN created_at TEXT",
        "ALTER TABLE usuarios ADD COLUMN updated_at TEXT",
    ]
    for comando in migracoes_usuarios:
        try:
            cursor.execute(comando)
        except sqlite3.OperationalError:
            pass  # coluna já existe — não faz nada

    agora = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    # Preenche created_at/updated_at só onde ainda estiver vazio (contas antigas
    # que não tinham essas colunas), sem sobrescrever contas já migradas.
    cursor.execute("UPDATE usuarios SET created_at = ? WHERE created_at IS NULL OR created_at = ''", (agora,))
    cursor.execute("UPDATE usuarios SET updated_at = ? WHERE updated_at IS NULL OR updated_at = ''", (agora,))

    # 2. A conta semente vira o primeiro Desenvolvedor (nível de acesso mais alto,
    #    acima de Admin — substitui o antigo papel "admin" que nunca chegou a ter
    #    UI própria). Contas que já eram 'admin' também sobem para 'desenvolvedor',
    #    sem perder nenhum dado.
    cursor.execute("UPDATE usuarios SET role = 'desenvolvedor' WHERE role = 'admin'")
    cursor.execute(
        "UPDATE usuarios SET role = 'desenvolvedor' WHERE LOWER(email) = 'samuel.ssousa1506@gmail.com' AND role != 'desenvolvedor'"
    )

    # 3. Migração de senhas em texto puro para hash seguro (werkzeug/pbkdf2).
    #    Identifica linhas cuja senha ainda NÃO está em formato de hash
    #    (hashes do werkzeug sempre começam com "pbkdf2:" ou "scrypt:") e
    #    converte em texto para hash, sem exigir que o professor troque a
    #    senha ou faça novo cadastro.
    cursor.execute("SELECT id, senha FROM usuarios")
    for usuario_id, senha_atual in cursor.fetchall():
        if senha_atual and not (senha_atual.startswith('pbkdf2:') or senha_atual.startswith('scrypt:')):
            cursor.execute(
                "UPDATE usuarios SET senha = ? WHERE id = ?",
                (generate_password_hash(senha_atual), usuario_id)
            )

    conn.commit()
    conn.close()

init_db()

def init_db_dev():
    """Tabelas da Central de Desenvolvimento — separado de init_db() para deixar
    claro que são aditivas e específicas dessa área técnica, sem mexer no schema
    principal do sistema."""
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS system_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_email TEXT,
            nome_usuario TEXT,
            modulo TEXT,
            acao TEXT,
            detalhes TEXT,
            status TEXT DEFAULT 'sucesso',
            criado_em TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS system_errors (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            modulo TEXT,
            usuario_email TEXT,
            mensagem TEXT,
            nivel TEXT DEFAULT 'aviso',
            status TEXT DEFAULT 'aberto',
            criado_em TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db_dev()

def registrar_log(usuario_email, nome_usuario, modulo, acao, detalhes='', status='sucesso'):
    """Registra uma ação no histórico técnico (Logs do Sistema). Nunca deve
    derrubar a requisição principal caso o log falhe — por isso o try/except
    silencioso: perder um log é aceitável, quebrar a geração de um material
    para o professor não é."""
    try:
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO system_logs (usuario_email, nome_usuario, modulo, acao, detalhes, status, criado_em)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (usuario_email, nome_usuario, modulo, acao, detalhes, status, datetime.now().strftime('%d/%m/%Y %H:%M:%S')))
        conn.commit()
        conn.close()
    except Exception:
        pass

def registrar_erro(modulo, usuario_email, mensagem, nivel='aviso'):
    """Registra uma falha na Central de Erros. nivel: 'info' | 'aviso' | 'atencao' | 'critico'."""
    try:
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO system_errors (modulo, usuario_email, mensagem, nivel, status, criado_em)
            VALUES (?, ?, ?, ?, 'aberto', ?)
        ''', (modulo, usuario_email, mensagem, nivel, datetime.now().strftime('%d/%m/%Y %H:%M:%S')))
        conn.commit()
        conn.close()
    except Exception:
        pass

# =====================================================================
# CONTROLE DE ACESSO (RBAC) — decorators para proteger rotas
# =====================================================================
def requer_login(f):
    """Garante que a rota só é acessada por quem está logado. Substitui
    gradualmente as checagens `if not session.get('logged_in')` repetidas
    manualmente em cada rota; novas rotas devem usar este decorator."""
    @wraps(f)
    def decorada(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorada

def requer_desenvolvedor(f):
    """Protege as rotas da Central de Desenvolvimento no BACKEND (não apenas
    escondendo o menu na interface). Bloqueia com 403 qualquer usuário que não
    seja Desenvolvedor — inclusive um professor que tente acessar a URL
    diretamente. Desenvolvedor é o nível de acesso mais alto do sistema."""
    @wraps(f)
    def decorada(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        if session.get('user_role') != 'desenvolvedor':
            abort(403)
        return f(*args, **kwargs)
    return decorada

# =====================================================================
# ROTAS FLASK
# =====================================================================
@app.errorhandler(403)
def acesso_negado(e):
    return render_template(
        'erro_acesso.html',
        app_name="Professor IA",
        name=session.get('user_name', ''), school=session.get('user_school', '')
    ), 403

@app.route('/desenvolvedor')
@requer_desenvolvedor
def central_desenvolvimento():
    modulos = [
        {'nome': 'Dashboard Técnico', 'icone': 'fa-chart-line', 'descricao': 'Indicadores reais de uso, geração de materiais e status do sistema.', 'disponivel': True, 'url': url_for('dev_dashboard_tecnico')},
        {'nome': 'Laboratório', 'icone': 'fa-flask', 'descricao': 'Teste prompts e modelos de IA sem afetar dados reais dos professores.', 'disponivel': False, 'url': '#'},
        {'nome': 'Central de IA', 'icone': 'fa-robot', 'descricao': 'Status da API, métricas de uso e editor de prompts com versionamento.', 'disponivel': False, 'url': '#'},
        {'nome': 'Gerenciador de Funcionalidades', 'icone': 'fa-toggle-on', 'descricao': 'Ative ou desative módulos do sistema para todos os professores.', 'disponivel': False, 'url': '#'},
        {'nome': 'Logs do Sistema', 'icone': 'fa-scroll', 'descricao': 'Histórico de ações dos usuários, com filtros por data, módulo e status.', 'disponivel': True, 'url': url_for('dev_logs')},
        {'nome': 'Central de Erros', 'icone': 'fa-triangle-exclamation', 'descricao': 'Monitoramento de falhas do sistema, classificadas por gravidade.', 'disponivel': True, 'url': url_for('dev_erros')},
        {'nome': 'Backups', 'icone': 'fa-database', 'descricao': 'Backup manual do banco de dados e histórico de versões salvas.', 'disponivel': False, 'url': '#'},
        {'nome': 'Controle de Versões', 'icone': 'fa-box-archive', 'descricao': 'Histórico de versões do Professor IA e o que mudou em cada uma.', 'disponivel': False, 'url': '#'},
        {'nome': 'Feature Flags', 'icone': 'fa-flag', 'descricao': 'Libere funcionalidades novas de forma controlada, por perfil ou usuário.', 'disponivel': False, 'url': '#'},
        {'nome': 'Usuários e Permissões', 'icone': 'fa-users-gear', 'descricao': 'Gerencie contas de professores: ativar, bloquear, redefinir senha, papel.', 'disponivel': False, 'url': '#'},
        {'nome': 'Configurações Técnicas', 'icone': 'fa-gears', 'descricao': 'Parâmetros técnicos gerais da plataforma.', 'disponivel': False, 'url': '#'},
    ]
    return render_template(
        'central_dev.html', modulos=modulos,
        name=session.get('user_name', ''), school=session.get('user_school', '')
    )

def _parse_data_log(texto):
    """Converte 'dd/mm/AAAA HH:MM:SS' (formato usado em todo o sistema) para datetime.
    Retorna None se não conseguir — nunca derruba a página por causa de uma data mal formada."""
    if not texto:
        return None
    for fmt in ('%d/%m/%Y %H:%M:%S', '%d/%m/%Y %H:%M'):
        try:
            return datetime.strptime(texto, fmt)
        except ValueError:
            continue
    return None

@app.route('/desenvolvedor/dashboard-tecnico')
@requer_desenvolvedor
def dev_dashboard_tecnico():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("SELECT COUNT(*) FROM usuarios")
    total_usuarios = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM usuarios WHERE ativo = 1")
    usuarios_ativos = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(DISTINCT escola) FROM usuarios WHERE escola != ''")
    total_escolas = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM usuarios WHERE role = 'professor' AND ativo = 1")
    professores_ativos = cursor.fetchone()[0]

    # Materiais gerados por módulo (dados reais — nada fictício)
    cursor.execute("SELECT tipo_modulo, COUNT(*) as qtd FROM materiais GROUP BY tipo_modulo")
    contagem_modulo = {row['tipo_modulo']: row['qtd'] for row in cursor.fetchall()}
    cursor.execute("SELECT COUNT(*) FROM materiais")
    total_materiais_gerados = cursor.fetchone()[0]

    # Erros nas últimas 24h — o campo criado_em é texto (dd/mm/AAAA HH:MM:SS), então
    # filtramos em Python em vez de confiar em comparação de string no SQLite.
    cursor.execute("SELECT criado_em, nivel FROM system_errors")
    agora = datetime.now()
    erros_24h = 0
    for row in cursor.fetchall():
        dt = _parse_data_log(row['criado_em'])
        if dt and (agora - dt) <= timedelta(hours=24):
            erros_24h += 1

    # Uso dos últimos 7 dias, por módulo — para o gráfico
    cursor.execute("SELECT tipo_modulo, criado_em FROM materiais")
    todos_materiais = cursor.fetchall()
    conn.close()

    dias = [(agora - timedelta(days=i)).date() for i in range(6, -1, -1)]
    uso_por_dia = {d: 0 for d in dias}
    for row in todos_materiais:
        dt = _parse_data_log(row['criado_em'])
        if dt and dt.date() in uso_por_dia:
            uso_por_dia[dt.date()] += 1
    maximo_uso = max(uso_por_dia.values()) if uso_por_dia and max(uso_por_dia.values()) > 0 else 1
    grafico_uso = [
        {
            'label': d.strftime('%d/%m'),
            'valor': uso_por_dia[d],
            'altura_pct': round((uso_por_dia[d] / maximo_uso) * 100)
        }
        for d in dias
    ]

    # Status do banco de dados — checagem real, não suposição
    try:
        conn_teste = sqlite3.connect('database.db')
        conn_teste.execute("SELECT 1")
        conn_teste.close()
        status_banco = True
    except Exception:
        status_banco = False

    return render_template(
        'dev_dashboard.html',
        name=session.get('user_name', ''),
        total_usuarios=total_usuarios, usuarios_ativos=usuarios_ativos,
        total_escolas=total_escolas, professores_ativos=professores_ativos,
        total_materiais_gerados=total_materiais_gerados,
        planejamentos=contagem_modulo.get('Plano de Aula', 0) + contagem_modulo.get('Planejamento Bimestral', 0),
        atividades=contagem_modulo.get('Banco de Atividades', 0),
        provas=contagem_modulo.get('Gerador de Provas', 0) + contagem_modulo.get('Simulados', 0),
        sequencias=contagem_modulo.get('Sequência Didática', 0),
        consultas_ia=contagem_modulo.get('Tira-Dúvidas com IA', 0),
        erros_24h=erros_24h,
        gemini_configurado=bool(GEMINI_API_KEY),
        mistral_configurado=bool(MISTRAL_API_KEY),
        status_banco=status_banco,
        grafico_uso=grafico_uso,
    )

@app.route('/desenvolvedor/logs')
@requer_desenvolvedor
def dev_logs():
    filtro_data = request.args.get('data', '').strip()
    filtro_usuario = request.args.get('usuario', '').strip()
    filtro_modulo = request.args.get('modulo', '').strip()
    filtro_status = request.args.get('status', '').strip()

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    condicoes = []
    parametros = []
    if filtro_usuario:
        condicoes.append("(usuario_email LIKE ? OR nome_usuario LIKE ?)")
        parametros.extend([f"%{filtro_usuario}%", f"%{filtro_usuario}%"])
    if filtro_modulo:
        condicoes.append("modulo = ?")
        parametros.append(filtro_modulo)
    if filtro_status:
        condicoes.append("status = ?")
        parametros.append(filtro_status)
    if filtro_data:
        try:
            data_fmt = datetime.strptime(filtro_data, '%Y-%m-%d').strftime('%d/%m/%Y')
            condicoes.append("substr(criado_em, 1, 10) = ?")
            parametros.append(data_fmt)
        except ValueError:
            pass

    where = f"WHERE {' AND '.join(condicoes)}" if condicoes else ""
    cursor.execute(f"SELECT * FROM system_logs {where} ORDER BY id DESC LIMIT 300", parametros)
    logs = cursor.fetchall()

    cursor.execute("SELECT DISTINCT modulo FROM system_logs WHERE modulo != '' ORDER BY modulo")
    modulos_disponiveis = [row['modulo'] for row in cursor.fetchall()]
    conn.close()

    return render_template(
        'dev_logs.html', logs=logs, modulos_disponiveis=modulos_disponiveis,
        filtro_data=filtro_data, filtro_usuario=filtro_usuario, filtro_modulo=filtro_modulo, filtro_status=filtro_status,
        name=session.get('user_name', '')
    )

@app.route('/desenvolvedor/erros')
@requer_desenvolvedor
def dev_erros():
    filtro_nivel = request.args.get('nivel', '').strip()
    filtro_status = request.args.get('status', '').strip()

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    condicoes = []
    parametros = []
    if filtro_nivel:
        condicoes.append("nivel = ?")
        parametros.append(filtro_nivel)
    if filtro_status:
        condicoes.append("status = ?")
        parametros.append(filtro_status)

    where = f"WHERE {' AND '.join(condicoes)}" if condicoes else ""
    cursor.execute(f"SELECT * FROM system_errors {where} ORDER BY id DESC LIMIT 300", parametros)
    erros = cursor.fetchall()
    conn.close()

    return render_template(
        'dev_erros.html', erros=erros, filtro_nivel=filtro_nivel, filtro_status=filtro_status,
        name=session.get('user_name', '')
    )

@app.route('/desenvolvedor/erros/<int:erro_id>/resolver', methods=['POST'])
@requer_desenvolvedor
def dev_resolver_erro(erro_id):
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("UPDATE system_errors SET status = 'resolvido' WHERE id = ?", (erro_id,))
    conn.commit()
    conn.close()
    registrar_log(session.get('user_email', ''), session.get('user_name', ''), 'Central de Erros', f'Marcou erro #{erro_id} como resolvido')
    return redirect(url_for('dev_erros'))

@app.route('/')
def index():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    return redirect(url_for('inicio'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    erro = None
    sucesso = request.args.get('sucesso')
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('senha', '').strip()

        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute("SELECT nome, escola, senha, role, ativo FROM usuarios WHERE LOWER(email) = ?", (email,))
        user = cursor.fetchone()

        if user and check_password_hash(user[2], password):
            if not user[4]:  # ativo == 0
                conn.close()
                registrar_log(email, user[0], 'Autenticação', 'Tentativa de login', 'Conta desativada', status='erro')
                erro = "Esta conta foi desativada. Entre em contato com o administrador."
                return render_template('login.html', erro=erro, sucesso=sucesso)

            session.permanent = True
            session['logged_in'] = True
            session['user_email'] = email
            session['user_name'] = user[0]
            session['user_school'] = user[1]
            session['user_role'] = user[3] or 'professor'
            conn.close()
            registrar_log(email, user[0], 'Autenticação', 'Login', status='sucesso')
            return redirect(url_for('inicio'))
        else:
            conn.close()
            registrar_log(email, '', 'Autenticação', 'Tentativa de login', 'Credenciais inválidas', status='erro')
            erro = "E-mail ou senha incorretos."
    return render_template('login.html', erro=erro, sucesso=sucesso)

@app.route('/cadastro', methods=['GET', 'POST'])
def cadastro():
    erro = None
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        escola = request.form.get('escola', '').strip()
        email = request.form.get('email', '').strip().lower()
        senha = request.form.get('senha', '').strip()
        
        if not nome or not escola or not email or not senha:
            erro = "Todos os campos são obrigatórios."
        else:
            try:
                agora = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
                conn = sqlite3.connect('database.db')
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO usuarios (nome, escola, email, senha, role, ativo, created_at, updated_at) VALUES (?, ?, ?, ?, 'professor', 1, ?, ?)",
                    (nome, escola, email, generate_password_hash(senha), agora, agora)
                )
                conn.commit()
                conn.close()
                return redirect(url_for('login', sucesso="Conta criada com sucesso! Faça login."))
            except sqlite3.IntegrityError:
                erro = "Este e-mail já está cadastrado no sistema."
    return render_template('cadastro.html', erro=erro)

@app.route('/inicio')
def inicio():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    email = session.get('user_email', '')
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Total de materiais e favoritos (dados reais do banco, sem números fictícios)
    cursor.execute("SELECT COUNT(*) FROM materiais WHERE usuario_email = ?", (email,))
    total_materiais = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM materiais WHERE usuario_email = ? AND favorito = 1", (email,))
    total_favoritos = cursor.fetchone()[0]

    # Contagem por módulo (para os cards de atalho mostrarem quantos materiais existem em cada um)
    cursor.execute("SELECT tipo_modulo, COUNT(*) as qtd FROM materiais WHERE usuario_email = ? GROUP BY tipo_modulo", (email,))
    contagem_por_modulo = {row['tipo_modulo']: row['qtd'] for row in cursor.fetchall()}

    # 5 materiais mais recentes
    cursor.execute("SELECT * FROM materiais WHERE usuario_email = ? ORDER BY id DESC LIMIT 5", (email,))
    materiais_recentes = cursor.fetchall()

    # 5 favoritos mais recentes
    cursor.execute("SELECT * FROM materiais WHERE usuario_email = ? AND favorito = 1 ORDER BY id DESC LIMIT 5", (email,))
    favoritos_recentes = cursor.fetchall()

    # "Continuar de onde parei" — o material mais recente, se existir algum
    continuar = materiais_recentes[0] if materiais_recentes else None

    conn.close()

    # Módulos existentes hoje, na ordem do menu, com a contagem real de cada um
    atalhos_ativos = []
    for chave, cfg in MODULOS.items():
        if chave in ('duvidas', 'relatorios'):
            continue  # não geram "materiais" salvos da mesma forma; ficam só no menu lateral
        atalhos_ativos.append({
            'chave': chave, 'nome': cfg['nome'], 'icone': cfg['icone'],
            'qtd': contagem_por_modulo.get(cfg['nome'], 0)
        })

    # Módulos da v2.0 ainda não implementados — aparecem no dashboard como prévia,
    # desabilitados, para não gerar um link quebrado.
    atalhos_em_breve = [
        {'nome': 'Diagnóstico da Turma', 'icone': 'fa-chart-pie'},
        {'nome': 'Assistente Pedagógico', 'icone': 'fa-comments'},
    ]

    hora_atual = datetime.now().hour
    if hora_atual < 12:
        saudacao = "Bom dia"
    elif hora_atual < 18:
        saudacao = "Boa tarde"
    else:
        saudacao = "Boa noite"

    return render_template(
        'inicio.html',
        app_name="Professor IA", name=session.get('user_name', ''), school=session.get('user_school', ''),
        saudacao=saudacao,
        total_materiais=total_materiais, total_favoritos=total_favoritos,
        materiais_recentes=materiais_recentes, favoritos_recentes=favoritos_recentes,
        continuar=continuar,
        atalhos_ativos=atalhos_ativos, atalhos_em_breve=atalhos_em_breve,
    )

@app.route('/dashboard', methods=['GET', 'POST'])
@app.route('/gerador', methods=['GET', 'POST'])
def dashboard():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    form_type = request.args.get('form_type') or request.form.get('form_type') or 'plano'
    if form_type not in MODULOS:
        form_type = 'plano'
        
    config_modulo = MODULOS[form_type]
    conteudo = ""
    
    tema = request.form.get('tema', '').strip()
    disciplina = request.form.get('disciplina', '').strip()
    ano = request.form.get('ano', '').strip()
    bncc = request.form.get('bncc', '').strip()
    tipo_prova = request.form.get('tipo_prova', '').strip()
    qtd_questoes = request.form.get('qtd_questoes', '').strip()
    nivel = request.form.get('nivel', '').strip()

    # Campos específicos do Planejamento Bimestral
    numero_plano = request.form.get('numero_plano', '').strip()
    bimestre = request.form.get('bimestre', '1º BIM').strip()
    data_inicio = request.form.get('data_inicio', '').strip()
    data_fim = request.form.get('data_fim', '').strip()
    turma = request.form.get('turma', '').strip()
    turno = request.form.get('turno', '').strip()
    modalidade = request.form.get('modalidade', 'Presencial').strip()
    ano_letivo = request.form.get('ano_letivo', '').strip()
    inep = request.form.get('inep', '').strip()
    endereco_escola = request.form.get('endereco_escola', '').strip()
    cidade_escola = request.form.get('cidade_escola', '').strip()
    estado_escola = request.form.get('estado_escola', 'MA').strip()
    zona_escola = request.form.get('zona_escola', '').strip()
    telefone_escola = request.form.get('telefone_escola', '').strip()
    email_escola = request.form.get('email_escola', '').strip()
    observacoes = request.form.get('observacoes', '').strip()

    # Campos específicos do Banco de Atividades e Projetos
    tipos_atividade = request.form.getlist('tipos_atividade')
    tipos_projeto = request.form.getlist('tipos_projeto')
    duracao = request.form.get('duracao', '').strip()
    objetivo = request.form.get('objetivo', '').strip()

    # Campos específicos do módulo de Alfabetização e Reforço de Leitura
    nome_aluno = request.form.get('nome_aluno', '').strip()
    nivel_leitura = request.form.get('nivel_leitura', '').strip()
    dificuldades_observadas = request.form.get('dificuldades_observadas', '').strip()
    focos_alfabetizacao = request.form.getlist('focos_alfabetizacao')
    duracao_alfabetizacao = request.form.get('duracao_alfabetizacao', '').strip()

    # Campos específicos do módulo de Simulados
    tipo_simulado = request.form.get('tipo_simulado', 'Simulado Geral / Diagnóstico').strip()
    duracao_simulado = request.form.get('duracao_simulado', '').strip()
    qtd_questoes_simulado = request.form.get('qtd_questoes_simulado', '20').strip()

    # Campos específicos da Sequência Didática
    quantidade_aulas = request.form.get('quantidade_aulas', '5').strip()
    objetivos_especificos = request.form.get('objetivos_especificos', '').strip()
    perfil_turma = request.form.get('perfil_turma', '').strip()
    dificuldades_turma = request.form.get('dificuldades_turma', '').strip()
    recursos_disponiveis = request.form.get('recursos_disponiveis', '').strip()
    metodologia_preferida = request.form.get('metodologia_preferida', '').strip()

    material_id = None
    info_pedagogica = ""

    pode_gerar = tema or (form_type == 'alfabetizacao' and (nivel_leitura or dificuldades_observadas or nome_aluno))
    if request.method == 'POST' and pode_gerar:
        conteudo, info_pedagogica, geracao_sucesso = executar_geracao_ia(
            tipo_modulo=config_modulo['nome'],
            disciplina=disciplina,
            ano=ano,
            tema=tema,
            bncc=bncc,
            tipo_prova=tipo_prova,
            qtd_questoes=qtd_questoes,
            nivel=nivel,
            nome_professor=session.get('user_name', 'Professor(a)'),
            nome_escola=session.get('user_school', 'Instituição de Ensino'),
            usuario_email=session.get('user_email', ''),
            numero_plano=numero_plano,
            bimestre=bimestre,
            data_inicio=data_inicio,
            data_fim=data_fim,
            turma=turma,
            turno=turno,
            modalidade=modalidade,
            ano_letivo=ano_letivo,
            inep=inep,
            endereco_escola=endereco_escola,
            cidade_escola=cidade_escola,
            estado_escola=estado_escola,
            zona_escola=zona_escola,
            telefone_escola=telefone_escola,
            email_escola=email_escola,
            observacoes=observacoes,
            tipos_atividade=tipos_atividade,
            tipos_projeto=tipos_projeto,
            duracao=duracao,
            objetivo=objetivo,
            nome_aluno=nome_aluno,
            nivel_leitura=nivel_leitura,
            dificuldades_observadas=dificuldades_observadas,
            focos_alfabetizacao=focos_alfabetizacao,
            duracao_alfabetizacao=duracao_alfabetizacao,
            tipo_simulado=tipo_simulado,
            duracao_simulado=duracao_simulado,
            qtd_questoes_simulado=qtd_questoes_simulado,
            quantidade_aulas=quantidade_aulas,
            objetivos_especificos=objetivos_especificos,
            perfil_turma=perfil_turma,
            dificuldades_turma=dificuldades_turma,
            recursos_disponiveis=recursos_disponiveis,
            metodologia_preferida=metodologia_preferida
        )

        registrar_log(
            session.get('user_email', ''), session.get('user_name', ''),
            config_modulo['nome'], f"Gerou {config_modulo['nome'].lower()}",
            detalhes=f"Tema: {tema[:120]}" if tema else '',
            status='sucesso' if geracao_sucesso else 'erro'
        )

        # Registra automaticamente no Histórico / Biblioteca
        try:
            conn = sqlite3.connect('database.db')
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO materiais (usuario_email, tipo_modulo, titulo, disciplina, ano, conteudo_html, favorito, criado_em, info_pedagogica)
                VALUES (?, ?, ?, ?, ?, ?, 0, ?, ?)
            ''', (
                session.get('user_email', ''), config_modulo['nome'], tema, disciplina, ano,
                conteudo, datetime.now().strftime('%d/%m/%Y %H:%M'), info_pedagogica
            ))
            material_id = cursor.lastrowid
            conn.commit()
            conn.close()
        except Exception:
            material_id = None

    return render_template(
        'dashboard.html',
        form_type=form_type,
        config=config_modulo,
        conteudo=conteudo,
        material_id=material_id,
        info_pedagogica=info_pedagogica,
        tema=tema,
        disciplina=disciplina if disciplina else "Componente Curricular",
        ano=ano,
        bncc=bncc,
        tipo_prova=tipo_prova,
        qtd_questoes=qtd_questoes,
        nivel=nivel,
        numero_plano=numero_plano,
        bimestre=bimestre,
        data_inicio=data_inicio,
        data_fim=data_fim,
        turma=turma,
        turno=turno,
        modalidade=modalidade,
        ano_letivo=ano_letivo,
        inep=inep,
        endereco_escola=endereco_escola,
        cidade_escola=cidade_escola,
        estado_escola=estado_escola,
        zona_escola=zona_escola,
        telefone_escola=telefone_escola,
        email_escola=email_escola,
        observacoes=observacoes,
        tipos_atividade=tipos_atividade,
        tipos_projeto=tipos_projeto,
        duracao=duracao,
        objetivo=objetivo,
        nome_aluno=nome_aluno,
        nivel_leitura=nivel_leitura,
        dificuldades_observadas=dificuldades_observadas,
        focos_alfabetizacao=focos_alfabetizacao,
        duracao_alfabetizacao=duracao_alfabetizacao,
        tipo_simulado=tipo_simulado,
        duracao_simulado=duracao_simulado,
        qtd_questoes_simulado=qtd_questoes_simulado,
        TIPOS_ATIVIDADE=TIPOS_ATIVIDADE,
        TIPOS_PROJETO=TIPOS_PROJETO,
        NIVEIS_LEITURA=NIVEIS_LEITURA,
        FOCOS_ALFABETIZACAO=FOCOS_ALFABETIZACAO,
        TIPOS_SIMULADO=TIPOS_SIMULADO,
        app_name="Professor IA",
        name=session.get('user_name', 'Samuel Araújo Sousa'),     
        school=session.get('user_school', 'U.E. Prof. João Martins Neto')  
    )

@app.route('/exportar/<int:material_id>/<formato>')
def exportar_material(material_id, formato):
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM materiais WHERE id = ? AND usuario_email = ?", (material_id, session.get('user_email', '')))
    material = cursor.fetchone()
    conn.close()

    if not material:
        return redirect(url_for('dashboard'))

    titulo = material['titulo'] or material['tipo_modulo']
    escola = session.get('user_school', '')
    professor = session.get('user_name', '')
    nome_arquivo_base = re.sub(r'[^a-zA-Z0-9]+', '_', titulo)[:60] or 'documento'

    if formato == 'docx':
        buffer = gerar_docx(titulo, escola, professor, material['conteudo_html'], material['tipo_modulo'], material['disciplina'], material['ano'])
        return send_file(
            buffer, as_attachment=True, download_name=f"{nome_arquivo_base}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    elif formato == 'pdf':
        buffer = gerar_pdf(titulo, escola, professor, material['conteudo_html'], material['tipo_modulo'], material['disciplina'], material['ano'])
        return send_file(buffer, as_attachment=True, download_name=f"{nome_arquivo_base}.pdf", mimetype='application/pdf')
    else:
        return redirect(url_for('dashboard'))

@app.route('/biblioteca')
def biblioteca():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    email = session.get('user_email', '')
    apenas_favoritos = request.args.get('favoritos') == '1'
    busca = request.args.get('busca', '').strip()
    filtro_disciplina = request.args.get('disciplina', '').strip()
    filtro_ano = request.args.get('ano', '').strip()
    filtro_tipo = request.args.get('tipo', '').strip()
    filtro_data_inicio = request.args.get('data_inicio', '').strip()
    filtro_data_fim = request.args.get('data_fim', '').strip()
    filtro_pasta = request.args.get('pasta', '').strip()  # '' = todas, 'sem_pasta' = sem pasta, ou id da pasta

    condicoes = ["usuario_email = ?"]
    parametros = [email]

    if apenas_favoritos:
        condicoes.append("favorito = 1")
    if busca:
        # Pesquisa por título, tema/conteúdo (guardado no título) ou disciplina, num único campo
        condicoes.append("(titulo LIKE ? OR disciplina LIKE ?)")
        termo = f"%{busca}%"
        parametros.extend([termo, termo])
    if filtro_disciplina:
        condicoes.append("disciplina LIKE ?")
        parametros.append(f"%{filtro_disciplina}%")
    if filtro_ano:
        condicoes.append("ano LIKE ?")
        parametros.append(f"%{filtro_ano}%")
    if filtro_tipo:
        condicoes.append("tipo_modulo = ?")
        parametros.append(filtro_tipo)
    if filtro_data_inicio:
        # criado_em é salvo como texto 'dd/mm/AAAA HH:MM:SS' — comparação segura via data ISO auxiliar não existe,
        # então filtramos pelo prefixo de data já formatado enviado pelo formulário (AAAA-MM-DD -> convertido abaixo).
        try:
            data_ini_fmt = datetime.strptime(filtro_data_inicio, '%Y-%m-%d').strftime('%d/%m/%Y')
            condicoes.append("substr(criado_em, 1, 10) >= ?")
            parametros.append(data_ini_fmt)
        except ValueError:
            pass
    if filtro_data_fim:
        try:
            data_fim_fmt = datetime.strptime(filtro_data_fim, '%Y-%m-%d').strftime('%d/%m/%Y')
            condicoes.append("substr(criado_em, 1, 10) <= ?")
            parametros.append(data_fim_fmt)
        except ValueError:
            pass
    if filtro_pasta == 'sem_pasta':
        condicoes.append("pasta_id IS NULL")
    elif filtro_pasta.isdigit():
        condicoes.append("pasta_id = ?")
        parametros.append(int(filtro_pasta))

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    query = f"SELECT * FROM materiais WHERE {' AND '.join(condicoes)} ORDER BY id DESC"
    cursor.execute(query, parametros)
    materiais = cursor.fetchall()

    cursor.execute("SELECT * FROM pastas WHERE usuario_email = ? ORDER BY nome COLLATE NOCASE", (email,))
    pastas = cursor.fetchall()

    # Lista de disciplinas e tipos já usados pelo professor, para preencher os filtros com opções reais
    cursor.execute("SELECT DISTINCT disciplina FROM materiais WHERE usuario_email = ? AND disciplina != '' ORDER BY disciplina", (email,))
    disciplinas_disponiveis = [row['disciplina'] for row in cursor.fetchall()]
    cursor.execute("SELECT DISTINCT tipo_modulo FROM materiais WHERE usuario_email = ? AND tipo_modulo != '' ORDER BY tipo_modulo", (email,))
    tipos_disponiveis = [row['tipo_modulo'] for row in cursor.fetchall()]

    conn.close()

    return render_template(
        'biblioteca.html', materiais=materiais, apenas_favoritos=apenas_favoritos,
        busca=busca, filtro_disciplina=filtro_disciplina, filtro_ano=filtro_ano,
        filtro_tipo=filtro_tipo, filtro_data_inicio=filtro_data_inicio, filtro_data_fim=filtro_data_fim,
        filtro_pasta=filtro_pasta, pastas=pastas,
        disciplinas_disponiveis=disciplinas_disponiveis, tipos_disponiveis=tipos_disponiveis,
        app_name="Professor IA", name=session.get('user_name', ''), school=session.get('user_school', '')
    )

@app.route('/biblioteca/pastas/criar', methods=['POST'])
def criar_pasta():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    nome = request.form.get('nome', '').strip()
    if nome:
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO pastas (usuario_email, nome, criado_em) VALUES (?, ?, ?)",
            (session.get('user_email', ''), nome, datetime.now().strftime('%d/%m/%Y %H:%M:%S'))
        )
        conn.commit()
        conn.close()
    return redirect(url_for('biblioteca'))

@app.route('/biblioteca/pastas/excluir/<int:pasta_id>', methods=['POST'])
def excluir_pasta(pasta_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    email = session.get('user_email', '')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # Os materiais dentro da pasta NÃO são excluídos — só voltam a ficar "sem pasta".
    cursor.execute("UPDATE materiais SET pasta_id = NULL WHERE pasta_id = ? AND usuario_email = ?", (pasta_id, email))
    cursor.execute("DELETE FROM pastas WHERE id = ? AND usuario_email = ?", (pasta_id, email))
    conn.commit()
    conn.close()
    return redirect(url_for('biblioteca'))

@app.route('/biblioteca/material/<int:material_id>/mover', methods=['POST'])
def mover_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    email = session.get('user_email', '')
    pasta_id = request.form.get('pasta_id', '').strip()

    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    if pasta_id and pasta_id.isdigit():
        # Confere que a pasta pertence mesmo a este professor antes de mover
        cursor.execute("SELECT id FROM pastas WHERE id = ? AND usuario_email = ?", (int(pasta_id), email))
        if cursor.fetchone():
            cursor.execute("UPDATE materiais SET pasta_id = ? WHERE id = ? AND usuario_email = ?", (int(pasta_id), material_id, email))
    else:
        cursor.execute("UPDATE materiais SET pasta_id = NULL WHERE id = ? AND usuario_email = ?", (material_id, email))
    conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('biblioteca'))

@app.route('/biblioteca/duplicar/<int:material_id>', methods=['POST'])
def duplicar_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    email = session.get('user_email', '')

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM materiais WHERE id = ? AND usuario_email = ?", (material_id, email))
    original = cursor.fetchone()

    if original:
        cursor.execute('''
            INSERT INTO materiais (usuario_email, tipo_modulo, titulo, disciplina, ano, conteudo_html, info_pedagogica, favorito, criado_em, pasta_id)
            VALUES (?, ?, ?, ?, ?, ?, ?, 0, ?, ?)
        ''', (
            email, original['tipo_modulo'], f"{original['titulo']} (Cópia)",
            original['disciplina'], original['ano'], original['conteudo_html'],
            original['info_pedagogica'], datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            original['pasta_id']
        ))
        novo_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return redirect(url_for('ver_material', material_id=novo_id))

    conn.close()
    return redirect(url_for('biblioteca'))

@app.route('/biblioteca/ver/<int:material_id>')
def ver_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM materiais WHERE id = ? AND usuario_email = ?", (material_id, session.get('user_email', '')))
    material = cursor.fetchone()
    conn.close()

    if not material:
        return redirect(url_for('biblioteca'))

    return render_template(
        'material.html', material=material,
        app_name="Professor IA", name=session.get('user_name', ''), school=session.get('user_school', '')
    )

@app.route('/biblioteca/favoritar/<int:material_id>', methods=['POST'])
def favoritar_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT favorito FROM materiais WHERE id = ? AND usuario_email = ?", (material_id, session.get('user_email', '')))
    row = cursor.fetchone()
    if row:
        novo_valor = 0 if row[0] == 1 else 1
        cursor.execute("UPDATE materiais SET favorito = ? WHERE id = ?", (novo_valor, material_id))
        conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('biblioteca'))

@app.route('/biblioteca/editar/<int:material_id>', methods=['POST'])
def editar_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    novo_titulo = request.form.get('title', '').strip()
    novo_conteudo = sanitizar_html_seguro(request.form.get('content', '').strip())
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE materiais SET titulo = ?, conteudo_html = ? WHERE id = ? AND usuario_email = ?",
        (novo_titulo, novo_conteudo, material_id, session.get('user_email', ''))
    )
    conn.commit()
    conn.close()

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return {'status': 'ok', 'material_id': material_id}

    return redirect(url_for('ver_material', material_id=material_id))

@app.route('/biblioteca/excluir/<int:material_id>', methods=['POST'])
def excluir_material(material_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM materiais WHERE id = ? AND usuario_email = ?", (material_id, session.get('user_email', '')))
    conn.commit()
    conn.close()
    return redirect(url_for('biblioteca'))

@app.route('/banco-questoes')
def banco_questoes():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    filtro_ano = request.args.get('ano', '').strip()
    filtro_disciplina = request.args.get('disciplina', '').strip()
    filtro_conteudo = request.args.get('conteudo', '').strip()
    filtro_bncc = request.args.get('bncc', '').strip()
    filtro_dificuldade = request.args.get('dificuldade', '').strip()

    query = "SELECT * FROM banco_questoes WHERE usuario_email = ?"
    params = [session.get('user_email', '')]
    if filtro_ano:
        query += " AND ano LIKE ?"; params.append(f"%{filtro_ano}%")
    if filtro_disciplina:
        query += " AND disciplina LIKE ?"; params.append(f"%{filtro_disciplina}%")
    if filtro_conteudo:
        query += " AND conteudo LIKE ?"; params.append(f"%{filtro_conteudo}%")
    if filtro_bncc:
        query += " AND bncc LIKE ?"; params.append(f"%{filtro_bncc}%")
    if filtro_dificuldade:
        query += " AND dificuldade = ?"; params.append(filtro_dificuldade)
    query += " ORDER BY id DESC"

    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute(query, params)
    questoes = cursor.fetchall()
    conn.close()

    return render_template(
        'banco_questoes.html', questoes=questoes,
        filtro_ano=filtro_ano, filtro_disciplina=filtro_disciplina, filtro_conteudo=filtro_conteudo,
        filtro_bncc=filtro_bncc, filtro_dificuldade=filtro_dificuldade,
        app_name="Professor IA", name=session.get('user_name', ''), school=session.get('user_school', '')
    )

@app.route('/banco-questoes/salvar', methods=['POST'])
def salvar_banco_questoes():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO banco_questoes (usuario_email, ano, disciplina, conteudo, bncc, dificuldade, enunciado_html, criado_em)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        session.get('user_email', ''),
        request.form.get('ano', ''),
        request.form.get('disciplina', ''),
        request.form.get('tema', ''),
        request.form.get('bncc', ''),
        request.form.get('nivel', ''),
        request.form.get('conteudo_html', ''),
        datetime.now().strftime('%d/%m/%Y %H:%M')
    ))
    conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('banco_questoes'))

@app.route('/banco-questoes/excluir/<int:questao_id>', methods=['POST'])
def excluir_questao(questao_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM banco_questoes WHERE id = ? AND usuario_email = ?", (questao_id, session.get('user_email', '')))
    conn.commit()
    conn.close()
    return redirect(url_for('banco_questoes'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)